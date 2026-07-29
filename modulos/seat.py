"""
modulos/seat.py - SEAT: Edicao e Triagem
Secretaria das Sessoes - TCDF

Sub-etapa 1A+1B: Pauta Ativa com sessoes + Distribuicao Equalitaria

Correcoes aplicadas:
- Delimitador CSV auto-detectado (; ou ,)
- Remocao do sufixo -e do numero do processo
- Adicao do prefixo GC nas iniciais do relator (exceto GAVF / Subst.)

Fluxo de status:
  inclusao -> em_edicao -> em_revisao -> encaminhado

RBAC:
- modo_edicao=True: pode incluir, distribuir, alterar status, editar atribuicoes
- modo_edicao=False: somente visualizacao
"""

import streamlit as st
import csv
import io
import unicodedata
from datetime import datetime, date
import db_manager
from modulos.gerenciar_dados import _renderizar_gerenciar_dados

try:
    import pandas as pd
    PANDAS_OK = True
except ImportError:
    PANDAS_OK = False

# ============================================================
# CONSTANTES
# ============================================================

STATUS_FLOW = {
    "inclusao": {
        "label": "Inclusao",
        "proximo": "em_edicao",
        "acao_proximo": "Iniciar Edicao",
    },
    "em_edicao": {
        "label": "Em Edicao",
        "proximo": "em_revisao",
        "acao_proximo": "Enviar para Revisao",
    },
    "em_revisao": {
        "label": "Em Revisao",
        "proximo": "encaminhado",
        "acao_proximo": "Encaminhar para SEXP",
    },
    "encaminhado": {
        "label": "Encaminhado",
        "proximo": None,
        "acao_proximo": None,
    },
}

TIPOS_SESSAO = [
    "Ordinaria",
    "Ordinaria Virtual",
    "Reservada",
    "Administrativa",
    "Urgente",
]

# ============================================================
# FUNCOES AUXILIARES: NORMALIZACAO E HIGIENIZACAO
# ============================================================

def _normalizar_texto(texto: str) -> str:
    """
    Normaliza texto para comparacao:
    minusculas, sem acentos, sem espacos extras.
    """
    if not texto:
        return ""
    texto = str(texto).lower().strip()
    texto = unicodedata.normalize('NFKD', texto)
    texto = ''.join(c for c in texto if not unicodedata.combining(c))
    texto = ' '.join(texto.split())
    return texto

def _normalizar_tipo_sessao(tipo: str) -> str:
    """
    Normaliza o tipo de sessao do CSV para corresponder aos valores padrao.
    Aceita variacoes como 'ordinaria', 'Ordinaria', 'ORDINARIA'.
    """
    if not tipo:
        return ""
    tipo_norm = _normalizar_texto(tipo)
    for tipo_padrao in TIPOS_SESSAO:
        if _normalizar_texto(tipo_padrao) == tipo_norm:
            return tipo_padrao
    return tipo.strip()

def _higienizar_numero_processo(numero: str) -> str:
    """
    Limpa o numero do processo:
    - Remove o sufixo -e (processo eletronico) do final
    
    Exemplo: 00600-00007999/2022-63-e -> 00600-00007999/2022-63
    """
    if not numero:
        return numero
    numero = numero.strip()
    if numero.lower().endswith("-e"):
        numero = numero[:-2]
    return numero

def _higienizar_relator(relator: str) -> str:
    """
    Formata o nome do relator:
    - Adiciona o prefixo GC antes das iniciais
    - Excecao: GAVF / Subst. e variantes sao mantidos como estao
    
    Exemplos:
      AM -> GCAM
      GAVF / Subst. -> GAVF / Subst. (mantido)
      GAVF / Subst -> GAVF / Subst (mantido)
    """
    if not relator:
        return relator
    relator = relator.strip()
    
    # Excecao: se ja contem GAVF ou Subst, manter como esta
    relator_upper = relator.upper()
    if "GAVF" in relator_upper or "SUBST" in relator_upper:
        return relator
    
    # Se ja comeca com GC, nao duplicar
    if relator_upper.startswith("GC"):
        return relator
    
    # Adicionar prefixo GC
    return f"GC{relator}"

def _higienizar_colaborador(nome_digitado: str, nomes_oficiais: list) -> str:
    """
    Faz matching inteligente entre nome digitado e nome oficial da equipe.
    Tolerante a variacoes de escrita (acentos, maiusculas, espacos).
    """
    if not nome_digitado or not nomes_oficiais:
        return nome_digitado or ""

    alvo = _normalizar_texto(nome_digitado)

    # 1. Match exato normalizado
    for oficial in nomes_oficiais:
        if _normalizar_texto(oficial) == alvo:
            return oficial

    # 2. Match por primeiro nome + ultimo nome
    partes_alvo = alvo.split()
    for oficial in nomes_oficiais:
        partes_oficial = _normalizar_texto(oficial).split()
        if len(partes_alvo) >= 2 and len(partes_oficial) >= 2:
            if partes_alvo[0] == partes_oficial[0] and partes_alvo[-1] == partes_oficial[-1]:
                return oficial

    # 3. Match por primeiro nome
    for oficial in nomes_oficiais:
        partes_oficial = _normalizar_texto(oficial).split()
        if partes_oficial and partes_alvo and partes_oficial[0] == partes_alvo[0]:
            return oficial

    return nome_digitado

def _formatar_data(data_iso: str) -> str:
    """Converte data ISO para DD/MM/YYYY HH:MM."""
    if not data_iso:
        return "-"
    try:
        dt = datetime.fromisoformat(str(data_iso).replace("Z", "+00:00"))
        return dt.strftime("%d/%m/%Y %H:%M")
    except (ValueError, TypeError):
        return str(data_iso) if data_iso else "-"

def _formatar_data_curta(data_iso: str) -> str:
    """Converte data ISO para DD/MM/YYYY."""
    if not data_iso:
        return "-"
    try:
        dt = datetime.fromisoformat(str(data_iso).replace("Z", "+00:00"))
        return dt.strftime("%d/%m/%Y")
    except (ValueError, TypeError):
        return str(data_iso) if data_iso else "-"

# ============================================================
# FUNCOES AUXILIARES: EQUIPE E AFASTAMENTOS
# ============================================================

def _obter_equipe_seat() -> list:
    """
    Retorna lista com os Nomes de Guerra de TODOS os membros da SEAT (incluindo Gerentes),
    lendo diretamente da Fonte Única da Verdade (usuarios_acesso).
    """
    try:
        todos = db_manager.buscar_todos("usuarios_acesso", filtros={"ativo": True}) or []
        membros_seat = [u for u in todos if _normalizar_texto(u.get("setor", "")) == "seat"]
        
        nomes = []
        for m in membros_seat:
            ng = m.get("nome_guerra")
            if ng and str(ng).strip():
                nomes.append(str(ng).strip())
            elif m.get("nome"):
                nomes.append(str(m["nome"]).strip().split()[0])
                
        return sorted(list(set(nomes)))
    except Exception:
        return []

def _obter_afastados() -> list:
    """Retorna lista de nomes de membros atualmente afastados."""
    return db_manager.listar_nomes_afastados()

def _obter_disponiveis() -> list:
    """Retorna membros disponiveis para distribuicao: equipe ativa - afastados."""
    equipe = _obter_equipe_seat()
    afastados = _obter_afastados()

    afastados_norm = set(_normalizar_texto(a) for a in afastados)

    disponiveis = []
    for nome in equipe:
        if _normalizar_texto(nome) not in afastados_norm:
            disponiveis.append(nome)

    return disponiveis

# ============================================================
# FUNCOES AUXILIARES: DUPLICIDADE
# ============================================================

def _verificar_duplicidade(processo_numero: str, numero_sessao: str, dia_sessao: str) -> bool:
    """Verifica se um processo ja esta cadastrado na mesma sessao/dia."""
    if not processo_numero or not numero_sessao or not dia_sessao:
        return False

    resultados = db_manager.buscar_todos("pauta_seat", filtros={
        "processo_numero": processo_numero,
        "numero_sessao": numero_sessao,
        "dia_sessao": dia_sessao,
    })
    return len(resultados) > 0

# ============================================================
# FUNCOES AUXILIARES: DISTRIBUICAO
# ============================================================

def _contar_atribuicoes(nomes: list, campo: str) -> dict:
    """Conta quantos processos estao atribuidos a cada nome como editor ou revisor."""
    todos_processos = db_manager.buscar_todos("pauta_seat")
    contador = {nome: 0 for nome in nomes}

    for proc in todos_processos:
        nome_atribuido = proc.get(campo)
        if nome_atribuido and nome_atribuido in contador:
            contador[nome_atribuido] += 1

    return contador

def _distribuir_processos(processos: list, editores: list, revisores: list) -> dict:
    """
    Algoritmo de distribuicao equalitaria para SEAT.
    
    Diferenca do SEXP:
    - No SEAT, todo mundo revisa todo mundo.
    - Editor PODE ser o mesmo que revisor (sem restricao).
    - Apenas balanceia a carga igualmente entre os disponiveis.
    """
    if not processos or not editores or not revisores:
        return {}

    contador_editor = _contar_atribuicoes(editores, "editor")
    contador_revisor = _contar_atribuicoes(revisores, "revisor")

    atribuicoes = {}

    for processo in processos:
        proc_id = processo.get("id")
        if not proc_id:
            continue

        # Editor: membro com menos atribuicoes
        editor = min(editores, key=lambda n: contador_editor[n])
        contador_editor[editor] += 1

        # Revisor: membro com menos atribuicoes (pode ser o mesmo do editor)
        revisor = min(revisores, key=lambda n: contador_revisor[n])
        contador_revisor[revisor] += 1

        atribuicoes[proc_id] = (editor, revisor)

    return atribuicoes

# ============================================================
# FUNCOES AUXILIARES: PARSING CSV
# ============================================================

def _detectar_delimitador(texto: str) -> str:
    """
    Detecta o delimitador do CSV automaticamente.
    Suporta ; (padrao brasileiro) e , (padrao internacional).
    Usa csv.Sniffer primeiro, com fallback para verificacao manual.
    """
    primeira_linha = texto.split('\n')[0] if texto else ""
    
    # Tentar com Sniffer (mais robusto)
    try:
        dialect = csv.Sniffer().sniff(primeira_linha, delimiters=';,')
        return dialect.delimiter
    except Exception:
        pass
    
    # Fallback: verificacao manual
    count_ponto_virgula = primeira_linha.count(';')
    count_virgula = primeira_linha.count(',')
    
    if count_ponto_virgula > count_virgula:
        return ';'
    elif count_virgula > 0:
        return ','
    else:
        return ';'  # Default brasileiro

def _parse_csv(arquivo) -> list:
    """
    Faz parse de um arquivo CSV e retorna lista de processos higienizados.

    Correcoes:
    - Remove BOM (Byte Order Mark) de arquivos salvos no Windows
    - Mostra colunas detectadas para debug
    - Aceita variacoes de nomes de coluna
    """
    if arquivo is None:
        return []

    # Ler conteudo
    conteudo = arquivo.read()

    # Decodificar - usar utf-8-sig para remover BOM automaticamente
    try:
        texto = conteudo.decode('utf-8-sig')
    except (UnicodeDecodeError, AttributeError):
        try:
            texto = conteudo.decode('latin-1')
        except Exception:
            texto = str(conteudo)

    # Remover BOM manualmente como fallback
    texto = texto.lstrip('\ufeff')

    # Detectar delimitador
    delimiter = _detectar_delimitador(texto)

    # Parse
    reader = csv.DictReader(io.StringIO(texto), delimiter=delimiter)

    # Debug: mostrar colunas detectadas
    if reader.fieldnames:
        st.caption(f"Colunas detectadas: {', '.join(reader.fieldnames)} | Delimitador: '{delimiter}'")

    processos = []
    for row in reader:
        if not row:
            continue

        # Normalizar nomes de coluna (lowercase, sem espacos, sem BOM)
        row_norm = {}
        for k, v in row.items():
            if k and v:
                # Remover BOM e normalizar nome da coluna
                chave = k.replace('\ufeff', '').lower().strip()
                row_norm[chave] = v.strip()

        # Buscar coluna de processo (aceita variacoes)
        processo_numero = (
            row_norm.get('processo_numero')
            or row_norm.get('processo')
            or row_norm.get('numero')
            or row_norm.get('n_processo')
            or row_norm.get('numero_processo')
            or row_norm.get('n_processo')
            or ""
        )

        # Buscar coluna de relator (opcional)
        relator = (
            row_norm.get('relator')
            or row_norm.get('relatora')
            or row_norm.get('relator(a)')
            or row_norm.get('rel')
            or None
        )

        # Buscar coluna de tipo de sessao
        tipo_sessao = (
            row_norm.get('tipo_sessao')
            or row_norm.get('tipo')
            or row_norm.get('sessao')
            or row_norm.get('tipo_sessao')
            or row_norm.get('tipo_ses')
            or ""
        )

        if processo_numero and tipo_sessao:
            processos.append({
                'processo_numero': _higienizar_numero_processo(processo_numero),
                'relator': _higienizar_relator(relator) if relator else None,
                'tipo_sessao': _normalizar_tipo_sessao(tipo_sessao),
            })

    return processos

# ============================================================
# TAB 1: PAUTA ATIVA
# ============================================================

def _incluir_processo_manual(modo_edicao: bool):
    """Formulario para incluir processo manualmente, um por vez."""
    st.markdown("### Incluir Processo")
    with st.form("form_incluir_manual"):
        col1, col2 = st.columns(2)
        with col1:
            processo_numero = st.text_input(
                "Numero do Processo *",
                placeholder="Ex: 00600-00007999/2022-63-e",
            )
            numero_sessao = st.text_input(
                "Numero da Sessao *",
                placeholder="Ex: 123/2026",
            )
        with col2:
            tipo_sessao = st.selectbox(
                "Tipo de Sessao *",
                options=TIPOS_SESSAO,
                index=0,
            )
            dia_sessao = st.date_input(
                "Dia da Sessao *",
                value=date.today(),
            )
        col3, col4 = st.columns(2)
        with col3:
            relator = st.text_input(
                "Relator (opcional)",
                placeholder="Ex: AM ou GAVF / Subst.",
            )
        with col4:
            observacoes = st.text_input(
                "Observacoes (opcional)",
                placeholder="Info adicional",
            )
        submit = st.form_submit_button("Incluir na Pauta", use_container_width=True)
        if submit:
            if not processo_numero.strip() or not numero_sessao.strip():
                st.error("Numero do processo e numero da sessao sao obrigatorios.")
                return
            # Higienizar dados antes de salvar
            numero_limpo = _higienizar_numero_processo(processo_numero.strip())
            relator_limpo = _higienizar_relator(relator.strip()) if relator else None
            dia_iso = dia_sessao.isoformat()
            # Verificar duplicidade
            if _verificar_duplicidade(numero_limpo, numero_sessao.strip(), dia_iso):
                st.error(
                    f"Duplicidade: o processo {numero_limpo} ja esta cadastrado "
                    f"na sessao {numero_sessao} de {_formatar_data_curta(dia_iso)}."
                )
                return
            dados = {
                "processo_numero": numero_limpo,
                "numero_sessao": numero_sessao.strip(),
                "dia_sessao": dia_iso,
                "tipo_sessao": tipo_sessao,
                "relator": relator_limpo,
                "status": "inclusao",
                "observacoes": observacoes.strip() if observacoes else "",
            }
            resultado = db_manager.inserir("pauta_seat", dados)
            if resultado:
                # GATILHO DS: identifica despacho singular pendente
                _identificar_ds_apos_inclusao(numero_limpo, resultado.get("id"))
                st.success(f"Processo {numero_limpo} incluido na pauta SEAT.")
                if relator and relator_limpo != relator.strip():
                    st.caption(f"Relator formatado: {relator.strip()} -> {relator_limpo}")
                if processo_numero.strip() != numero_limpo:
                    st.caption(f"Numero limpo: {processo_numero.strip()} -> {numero_limpo}")
                st.rerun()
            else:
                st.error("Erro ao incluir processo. Verifique a conexao com o banco.")

def _incluir_processo_lote(modo_edicao: bool):
    """Upload de arquivo CSV com multiplos processos."""
    st.markdown("### Incluir em Lote (CSV)")
    st.caption(
        "Formato esperado: colunas `processo_numero`, `relator` (opcional), `tipo_sessao`.\n\n"
        "Delimitador automatico: aceita `;` (padrao BR) ou `,`.\n\n"
        "O sistema remove o sufixo `-e` do numero do processo e adiciona `GC` "
        "antes das iniciais do relator (exceto GAVF / Subst.)."
    )
    arquivo = st.file_uploader(
        "Selecionar arquivo CSV",
        type=['csv'],
        key="csv_upload_seat",
    )
    if arquivo is None:
        return
    # Parse do CSV
    processos_csv = _parse_csv(arquivo)
    if not processos_csv:
        st.error(
            "Nenhum processo valido encontrado no arquivo. "
            "Verifique se as colunas `processo_numero` e `tipo_sessao` existem e estao preenchidas."
        )
        return
    # Detectar tipos de sessao unicos no arquivo
    tipos_encontrados = sorted(set(p['tipo_sessao'] for p in processos_csv if p['tipo_sessao']))
    st.success(
        f"CSV carregado: {len(processos_csv)} processos encontrados, "
        f"{len(tipos_encontrados)} tipo(s) de sessao."
    )
    # Mostrar preview dos dados higienizados
    with st.expander("Preview dos dados (apos higienizacao)", expanded=False):
        for p in processos_csv[:10]:
            st.write(
                f"- Processo: {p['processo_numero']} | "
                f"Relator: {p.get('relator', '-') or '-'} | "
                f"Tipo: {p['tipo_sessao']}"
            )
        if len(processos_csv) > 10:
            st.caption(f"... e mais {len(processos_csv) - 10} processo(s).")
    # Formulario para numero e data de cada tipo de sessao
    st.markdown("### Informacoes das Sessoes")
    st.markdown("Preencha o numero e a data para cada tipo de sessao encontrado no arquivo:")
    session_info = {}
    for i, tipo in enumerate(tipos_encontrados):
        st.markdown(f"**{tipo}**")
        col_n, col_d = st.columns(2)
        with col_n:
            numero = st.text_input(
                "Numero da Sessao *",
                placeholder=f"Ex: 123/2026",
                key=f"csv_sessao_num_{i}",
            )
        with col_d:
            dia = st.date_input(
                "Data da Sessao *",
                value=date.today(),
                key=f"csv_sessao_dia_{i}",
            )
        session_info[tipo] = {
            'numero': numero.strip(),
            'dia': dia.isoformat(),
        }
        st.markdown("---")
    # Botao de confirmacao
    if st.button("Confirmar e Inserir", type="primary", key="csv_confirmar"):
        erros_validacao = []
        for tipo, info in session_info.items():
            if not info['numero']:
                erros_validacao.append(f"Numero da sessao para {tipo} e obrigatorio.")
        if erros_validacao:
            for erro in erros_validacao:
                st.error(erro)
            return
        inseridos = 0
        duplicados = 0
        erros = 0
        lista_duplicados = []
        for proc in processos_csv:
            tipo = proc['tipo_sessao']
            info = session_info.get(tipo)
            if not info:
                erros += 1
                continue
            if _verificar_duplicidade(proc['processo_numero'], info['numero'], info['dia']):
                lista_duplicados.append(
                    f"{proc['processo_numero']} (sessao {info['numero']} de {_formatar_data_curta(info['dia'])})"
                )
                duplicados += 1
                continue
            dados = {
                "processo_numero": proc['processo_numero'],
                "numero_sessao": info['numero'],
                "dia_sessao": info['dia'],
                "tipo_sessao": tipo,
                "relator": proc.get('relator'),
                "status": "inclusao",
            }
            resultado = db_manager.inserir("pauta_seat", dados)
            if resultado:
                # GATILHO DS: identifica despacho singular pendente
                _identificar_ds_apos_inclusao(proc['processo_numero'], resultado.get("id"))
                inseridos += 1
            else:
                erros += 1
        st.success(
            f"Importacao concluida: {inseridos} inseridos, "
            f"{duplicados} duplicados, {erros} erros."
        )
        if lista_duplicados:
            with st.expander(f"Ver {len(lista_duplicados)} processo(s) duplicado(s)"):
                for dup in lista_duplicados:
                    st.warning(f"Duplicado: {dup}")
        if inseridos > 0:
            st.rerun()

def _avancar_status(id_processo: int, status_atual: str):
    """Avanca o status do processo para a proxima etapa."""
    info_status = STATUS_FLOW.get(status_atual)
    if not info_status or not info_status["proximo"]:
        return

    proximo_status = info_status["proximo"]
    dados_update = {"status": proximo_status}

    if proximo_status == "encaminhado":
        dados_update["data_conclusao"] = datetime.now().isoformat()

    resultado = db_manager.atualizar("pauta_seat", id_processo, dados_update)

    if resultado:
        st.success(f"Processo movido para: {STATUS_FLOW[proximo_status]['label']}")
        st.rerun()
    else:
        st.error("Erro ao atualizar status do processo.")

def _voltar_status(id_processo: int, status_atual: str):
    """Volta o status do processo para a etapa anterior."""
    ordem = ["inclusao", "em_edicao", "em_revisao", "encaminhado"]
    indice = ordem.index(status_atual)
    if indice > 0:
        status_anterior = ordem[indice - 1]
        dados_update = {"status": status_anterior}
        if status_atual == "encaminhado":
            dados_update["data_conclusao"] = None

        resultado = db_manager.atualizar("pauta_seat", id_processo, dados_update)
        if resultado:
            st.success(f"Processo retornado para: {STATUS_FLOW[status_anterior]['label']}")
            st.rerun()

def _remover_processo(id_processo: int, numero: str):
    """Remove um processo da pauta com confirmacao."""
    chave_confirm = f"confirmar_remocao_{id_processo}"
    if st.session_state.get(chave_confirm):
        resultado = db_manager.deletar("pauta_seat", id_processo)
        if resultado:
            st.success(f"Processo {numero} removido da pauta.")
            del st.session_state[chave_confirm]
            st.rerun()
        else:
            st.error("Erro ao remover processo.")
            del st.session_state[chave_confirm]
            st.rerun()
    else:
        st.session_state[chave_confirm] = True
        st.warning(f"Confirme a remocao do processo {numero} clicando novamente.")
        st.rerun()

def _marcar_editado(id_proc, valor):
    """Marca/desmarca editado e atualiza o status automaticamente."""
    if valor:
        # Marcando como editado → status em_edicao
        db_manager.atualizar("pauta_seat", id_proc, {
            "editado": True,
            "status": "em_edicao",
        })
    else:
        # Desmarcando editado → desmarca revisado também e volta status
        db_manager.atualizar("pauta_seat", id_proc, {
            "editado": False,
            "revisado": False,
            "status": "inclusao",
        })
    st.rerun()

def _marcar_revisado(id_proc, valor):
    """Marca/desmarca revisado e encaminha para SEXP automaticamente."""
    if valor:
        # Marcando como revisado → encaminhado automaticamente para SEXP
        from datetime import datetime
        db_manager.atualizar("pauta_seat", id_proc, {
            "revisado": True,
            "status": "encaminhado",
            "data_conclusao": datetime.now().isoformat(),
        })
    else:
        # Desmarcando revisado → volta para em_revisao
        db_manager.atualizar("pauta_seat", id_proc, {
            "revisado": False,
            "status": "em_revisao",
            "data_conclusao": None,
        })
    st.rerun()

def _marcar_revisado(id_processo: int, valor: bool):
    """Marca/desmarca o checkbox de revisado e atualiza o status automaticamente."""
    processo = db_manager.buscar_por_id("pauta_seat", id_processo)
    if not processo:
        return

    editado = processo.get("editado", False)

    if valor and editado:
        novo_status = "encaminhado"
    elif valor and not editado:
        novo_status = "encaminhado"
    elif not valor and editado:
        novo_status = "em_revisao"
    else:
        novo_status = "em_edicao"

    db_manager.atualizar("pauta_seat", id_processo, {
        "revisado": valor,
        "status": novo_status,
    })
    st.rerun()

def _salvar_comentario(id_processo: int, comentario: str):
    """Salva o comentario de um processo."""
    db_manager.atualizar("pauta_seat", id_processo, {
        "comentario": comentario.strip(),
    })
    st.success("Comentario salvo.")
    st.rerun()


def renderizar_sidebar(usuario: dict, modo_edicao: bool = False):
    """
    Renderiza tabelas de carga na barra lateral.
    - Filtra apenas as sessoes mais recentes de cada tipo
    - Exclui Urgente (rito diferente)
    - Mostra apenas membros que participaram da distribuicao
    - Operacionais: veem apenas seus proprios dados
    - Gerente e acima: veem todos os membros que participaram
    """
    import pandas as pd

    cargo_usuario = usuario.get("cargo", "operacional")
    nome_usuario = usuario.get("nome", "")
    filtrar_por_usuario = (cargo_usuario == "operacional" and nome_usuario)

    todos_processos = db_manager.buscar_todos("pauta_seat")
    if not todos_processos:
        return

    datas_recentes = {}
    for p in todos_processos:
        tipo = p.get("tipo_sessao", "")
        dia = p.get("dia_sessao")
        if not tipo or not dia:
            continue
        if "urgente" in _normalizar_texto(tipo):
            continue
        dia_str = str(dia)[:10]
        tipo_key = _normalizar_texto(tipo)
        if tipo_key not in datas_recentes or dia_str > datas_recentes[tipo_key]:
            datas_recentes[tipo_key] = dia_str

    if not datas_recentes:
        return

    processos_recentes = []
    for p in todos_processos:
        tipo = p.get("tipo_sessao", "")
        dia = p.get("dia_sessao")
        if not tipo or not dia:
            continue
        dia_str = str(dia)[:10]
        tipo_key = _normalizar_texto(tipo)
        if tipo_key in datas_recentes and dia_str == datas_recentes[tipo_key]:
            processos_recentes.append(p)

    if not processos_recentes:
        return

    nomes_participantes = set()
    for p in processos_recentes:
        editor = (p.get("editor") or "").strip()
        revisor = (p.get("revisor") or "").strip()
        if editor:
            nomes_participantes.add(_normalizar_texto(editor))
        if revisor:
            nomes_participantes.add(_normalizar_texto(revisor))

    if not nomes_participantes:
        return

    equipe = _obter_equipe_seat()
    equipe_participante = []
    for nome in equipe:
        if _normalizar_texto(nome) in nomes_participantes:
            equipe_participante.append(nome)

    if not equipe_participante:
        return

    if filtrar_por_usuario:
        nome_norm = _normalizar_texto(nome_usuario)
        equipe_filtrada = [n for n in equipe_participante if _normalizar_texto(n) == nome_norm]
    else:
        equipe_filtrada = equipe_participante

    dados_edicao = []
    dados_revisao = []
    for nome in equipe_filtrada:
        nome_norm = _normalizar_texto(nome)
        qtd_editar = 0
        faltam_editar = 0
        qtd_revisar = 0
        faltam_revisar = 0
        for p in processos_recentes:
            editor_p = _normalizar_texto(p.get("editor", "") or "")
            revisor_p = _normalizar_texto(p.get("revisor", "") or "")
            editado_p = bool(p.get("editado", False))
            revisado_p = bool(p.get("revisado", False))
            if editor_p == nome_norm:
                qtd_editar += 1
                if not editado_p:
                    faltam_editar += 1
            if revisor_p == nome_norm:
                qtd_revisar += 1
                if not revisado_p:
                    faltam_revisar += 1
        dados_edicao.append({"Resp.": nome, "Qtd": qtd_editar, "Faltam": faltam_editar})
        dados_revisao.append({"Resp.": nome, "Qtd": qtd_revisar, "Faltam": faltam_revisar})

    if not dados_edicao:
        return

    df_ed = pd.DataFrame(dados_edicao)
    df_rev = pd.DataFrame(dados_revisao)

    if not filtrar_por_usuario:
        df_ed = pd.concat([df_ed, pd.DataFrame([{"Resp.": "Total", "Qtd": df_ed["Qtd"].sum(), "Faltam": df_ed["Faltam"].sum()}])], ignore_index=True)
        df_rev = pd.concat([df_rev, pd.DataFrame([{"Resp.": "Total", "Qtd": df_rev["Qtd"].sum(), "Faltam": df_rev["Faltam"].sum()}])], ignore_index=True)

    # TUDO DENTRO DA SIDEBAR
    with st.sidebar:
        st.markdown("**Edicao**")
        st.dataframe(df_ed, hide_index=True, use_container_width=True, height=len(df_ed) * 35 + 40)
        st.markdown("**Revisao**")
        st.dataframe(df_rev, hide_index=True, use_container_width=True, height=len(df_rev) * 35 + 40)

def _renderizar_sidebar_ds(usuario: dict):
    """
    Mostra os ultimos 5 DS e 5 Sustentacoes Orais na sidebar.
    Apenas criador, raiz e gerente.
    """
    import pandas as pd

    cargo = usuario.get("cargo", "operacional")
    if cargo not in ("criador", "raiz", "gerente"):
        return

    todos_ds = db_manager.buscar_todos(
        "despachos_ds",
        ordem_coluna="created_at",
        ordem_desc=True,
    )

    if not todos_ds:
        return

    ds_lista = [d for d in todos_ds if d.get("tipo") == "Despacho Singular"][:5]
    so_lista = [d for d in todos_ds if d.get("tipo") == "Sustentacao Oral"][:5]

    with st.sidebar:
        st.markdown("---")

        if ds_lista:
            st.markdown("##### Despachos Singulares (Recentes)")
            dados_ds = []
            for ds in ds_lista:
                oficios = db_manager.buscar_todos(
                    "oficios_ds",
                    filtros={"despacho_id": ds["id"]},
                )
                dados_ds.append({
                    "Processo": ds.get("processo_numero", ""),
                    "Relator": ds.get("relator", "-") or "-",
                    "Docs": len(oficios),
                })
            df_ds = pd.DataFrame(dados_ds)
            st.dataframe(
                df_ds,
                hide_index=True,
                use_container_width=True,
                height=len(df_ds) * 35 + 40,
            )

        if so_lista:
            st.markdown("##### Sustentacao Oral (Recentes)")
            dados_so = []
            for so in so_lista:
                dados_so.append({
                    "Processo": so.get("processo_numero", ""),
                    "Relator": so.get("relator", "-") or "-",
                    "Confirmada": "Sim" if so.get("recebido_confirmado") else "Nao",
                })
            df_so = pd.DataFrame(dados_so)
            st.dataframe(
                df_so,
                hide_index=True,
                use_container_width=True,
                height=len(df_so) * 35 + 40,
            )

def _renderizar_card_processo(processo: dict, modo_edicao: bool):
    """Renderiza um card individual de processo na pauta ativa."""
    id_proc = processo.get("id")
    numero = processo.get("processo_numero", "")
    numero_sessao = processo.get("numero_sessao", "") or "-"
    dia_sessao = _formatar_data_curta(processo.get("dia_sessao"))
    tipo_sessao = processo.get("tipo_sessao", "") or "-"
    relator = processo.get("relator", "") or "-"
    editor = processo.get("editor", "") or "-"
    revisor = processo.get("revisor", "") or "-"
    editado = processo.get("editado", False)
    revisado = processo.get("revisado", False)
    status = processo.get("status", "inclusao")
    comentario = processo.get("comentario", "") or ""
    data_entrada = _formatar_data(processo.get("data_entrada"))

    # Ícone de status
    if status == "encaminhado":
        icone_status = "📤"
    elif revisado:
        icone_status = "✅"
    elif editado:
        icone_status = "📝"
    else:
        icone_status = "⏳"

    with st.container():
        # Linha 1: Processo + Relator + Status
        col_proc, col_rel, col_status = st.columns([3, 2, 1])
        with col_proc:
            st.markdown(f"### {icone_status} {numero}")
        with col_rel:
            st.markdown(f"**Relator:** {relator}")
        with col_status:
            if status == "encaminhado":
                st.success("Encaminhado")
            elif editado and not revisado:
                st.info("Em revisão")
            else:
                st.caption("Aguardando")

        # Linha 2: Sessao + Tipo + Data
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown(f"**Sessão:** {numero_sessao}")
        with col2:
            st.markdown(f"**Tipo:** {tipo_sessao}")
        with col3:
            st.markdown(f"**Data:** {dia_sessao}")

        # Linha 3: Editor + Revisor
        col1, col2 = st.columns(2)
        with col1:
            st.markdown(f"**Editor:** {editor}")
        with col2:
            st.markdown(f"**Revisor:** {revisor}")

        # Linha 4: Checkboxes Editado / Revisado
        if modo_edicao:
            col_chk1, col_chk2 = st.columns(2)
            with col_chk1:
                novo_editado = st.checkbox(
                    "Editado", value=editado, key=f"chk_editado_{id_proc}"
                )
                if novo_editado != editado:
                    _marcar_editado(id_proc, novo_editado)

            with col_chk2:
                # TRAVA: Revisado só funciona se Editado estiver marcado
                novo_revisado = st.checkbox(
                    "Revisado",
                    value=revisado,
                    key=f"chk_revisado_{id_proc}",
                    disabled=not novo_editado,  # ← TRAVA AQUI
                    help="Marque como editado primeiro" if not novo_editado else None,
                )
                if novo_revisado != revisado and novo_editado:
                    _marcar_revisado(id_proc, novo_revisado)
        else:
            col_chk1, col_chk2 = st.columns(2)
            with col_chk1:
                st.markdown(f"{'☑' if editado else '☐'} Editado")
            with col_chk2:
                st.markdown(f"{'☑' if revisado else '☐'} Revisado")

        # Linha 5: Comentario
        if modo_edicao:
            novo_comentario = st.text_area(
                "Comentário",
                value=comentario,
                placeholder="Deixe um comentário sobre o processo...",
                height=60,
                key=f"comentario_{id_proc}",
            )
            if st.button("Salvar Comentário", key=f"btn_comentario_{id_proc}"):
                if novo_comentario.strip() != comentario:
                    _salvar_comentario(id_proc, novo_comentario)
        else:
            if comentario:
                st.markdown(f"**Comentário:** {comentario}")
            else:
                st.caption("Sem comentário.")

        # Rodapé
        st.caption(f"Entrada: {data_entrada}")

        # Botão de remover
        if modo_edicao:
            if st.button("Remover", key=f"remover_{id_proc}"):
                _remover_processo(id_proc, numero)

        st.markdown("---")


def _renderizar_pauta_ativa(modo_edicao: bool, usuario: dict = None):
    """Renderiza a aba de Pauta Ativa com filtros e lista de processos."""
    cargo_usuario = usuario.get("cargo", "operacional") if usuario else "operacional"
    nome_usuario = usuario.get("nome", "") if usuario else ""
    filtrar_por_usuario = (cargo_usuario == "operacional" and nome_usuario)

    col_f1, col_f2, col_f3 = st.columns(3)

    with col_f1:
        busca = st.text_input(
            "Buscar por numero do processo",
            placeholder="Digite o numero...",
            key="busca_pauta_seat",
        )
    with col_f2:
        status_opcoes = ["todos"] + list(STATUS_FLOW.keys())
        status_labels = {"todos": "Todos os Status"}
        status_labels.update({k: v["label"] for k, v in STATUS_FLOW.items()})
        filtro_status = st.selectbox(
            "Filtrar por status",
            options=status_opcoes,
            format_func=lambda x: status_labels[x],
            key="filtro_status_seat",
        )
    with col_f3:
        filtro_tipo = st.selectbox(
            "Filtrar por tipo de sessao",
            options=["todos"] + TIPOS_SESSAO,
            key="filtro_tipo_seat",
        )

    if modo_edicao:
        tab_manual, tab_lote = st.tabs(["Inclusao Manual", "Inclusao em Lote (CSV)"])
        with tab_manual:
            _incluir_processo_manual(modo_edicao)
        with tab_lote:
            _incluir_processo_lote(modo_edicao)

    filtros = {}
    if filtro_status != "todos":
        filtros["status"] = filtro_status
    if filtro_tipo != "todos":
        filtros["tipo_sessao"] = filtro_tipo

    processos = db_manager.buscar_todos(
        "pauta_seat",
        filtros=filtros if filtros else None,
        ordem_coluna="created_at",
        ordem_desc=True,
    )

    # Filtro de busca por numero (client-side)
    if busca.strip():
        busca_lower = busca.strip().lower()
        processos = [p for p in processos if busca_lower in (p.get("processo_numero", "") or "").lower()]

    # FILTRAR POR USUARIO: operacionais so veem seus processos
    if filtrar_por_usuario:
        nome_norm = _normalizar_texto(nome_usuario)
        processos = [
            p for p in processos
            if _normalizar_texto(p.get("editor", "")) == nome_norm
            or _normalizar_texto(p.get("revisor", "")) == nome_norm
        ]

    # === Separar encaminhados dos ativos ===
    encaminhados = [p for p in processos if p.get("status") == "encaminhado"]
    ativos = [p for p in processos if p.get("status") != "encaminhado"]

    # Contadores
    col_c1, col_c2, col_c3, col_c4, col_c5 = st.columns(5)
    with col_c1:
        st.metric("Total", len(processos))
    with col_c2:
        st.metric("Inclusao", len([p for p in ativos if p.get("status") == "inclusao"]))
    with col_c3:
        st.metric("Em Edicao", len([p for p in ativos if p.get("status") == "em_edicao"]))
    with col_c4:
        st.metric("Em Revisao", len([p for p in ativos if p.get("status") == "em_revisao"]))
    with col_c5:
        st.metric("Encaminhados", len(encaminhados))

    st.markdown("---")

    # === Botão Finalizar Sessão (Liberado para toda a equipe operando na SEAT) ===
    if processos and modo_edicao:
        st.markdown("---")
        pendentes = len(processos) - len(encaminhados)

        if pendentes == 0:
            st.success(f"🎉 **Todos os {len(encaminhados)} processos foram revisados! A sessão está pronta para envio.**")
        else:
            st.warning(
                f"⚠️ **Alerta de Pendência:** Ainda há **{pendentes}** processo(s) sem revisão final nesta pauta.\n\n"
                f"Se você finalizar agora, apenas os **{len(encaminhados)}** processos já revisados serão enviados para a SEXP."
            )

        # Trava de segurança: se houver pendentes, exige marcação do checkbox para liberar o botão
        confirmar_envio = True
        if pendentes > 0:
            confirmar_envio = st.checkbox(
                f"Estou ciente das pendências e desejo finalizar a sessão enviando apenas os {len(encaminhados)} processos revisados para a SEXP.",
                key="chk_confirmar_finalizacao_seat"
            )

        if st.button(
            "📋 Finalizar Sessão e Enviar para SEXP",
            type="primary",
            use_container_width=True,
            disabled=not confirmar_envio,
            key="btn_finalizar_sessao_geral"
        ):
            if len(encaminhados) == 0:
                st.error("Nenhum processo foi revisado ainda. Impossível enviar sessão vazia para a SEXP.")
            else:
                with st.spinner("Enviando processos para a SEXP..."):
                    salvos = 0
                    ignorados = 0
                    for p in encaminhados:
                        # TRAVA OTIMISTA: Verifica se outro usuário já finalizou antes de regravar
                        proc_atual = db_manager.buscar_por_id("pauta_seat", p["id"])
                        if proc_atual and proc_atual.get("sessao_finalizada"):
                            ignorados += 1
                            continue

                        # Marca como finalizada para destravar a leitura no SEXP
                        res = db_manager.atualizar("pauta_seat", p["id"], {
                            "sessao_finalizada": True,
                            "status": "encaminhado"
                        })
                        if res:
                            salvos += 1

                if ignorados > 0:
                    st.warning(f"⚠️ {ignorados} processo(s) ignorado(s) pois já haviam sido enviados por outro gerente. Tela atualizada.")
                if salvos > 0:
                    st.success(f"✅ Sessão finalizada com sucesso! {salvos} processo(s) enviado(s) para a SEXP.")
                st.rerun()

        st.caption("Ao finalizar, os processos revisados migram para a esteira da Expedição (SEXP) e não poderão mais ser alterados na SEAT.")
        st.markdown("---")

    # === Listar processos ativos (nao encaminhados) ===
    if not ativos:
        if filtrar_por_usuario:
            st.info(f"Todos os seus processos foram encaminhados para o SEXP. ✅")
        else:
            st.info("Todos os processos foram encaminhados para o SEXP. ✅")
    else:
        if filtrar_por_usuario:
            st.markdown(f"### Meus Processos ({len(ativos)})")
        else:
            st.markdown(f"### Pauta Ativa ({len(ativos)} processo{'s' if len(ativos) != 1 else ''})")
        for processo in ativos:
            _renderizar_card_processo(processo, modo_edicao)

# ============================================================
# TAB 2: DISTRIBUICAO
# ============================================================

def _renderizar_distribuicao(modo_edicao: bool, usuario: dict = None):
    """Renderiza a aba de Distribuicao Equalitaria."""

    # Determinar se precisa filtrar por usuario
    cargo_usuario = usuario.get("cargo", "operacional") if usuario else "operacional"
    nome_usuario = usuario.get("nome", "") if usuario else ""
    filtrar_por_usuario = (cargo_usuario == "operacional" and nome_usuario)

    todos_processos = db_manager.buscar_todos(
        "pauta_seat",
        ordem_coluna="created_at",
        ordem_desc=True,
    )

    # FILTRAR POR USUARIO: operacionais so veem seus processos
    if filtrar_por_usuario:
        nome_norm = _normalizar_texto(nome_usuario)
        todos_processos = [
            p for p in todos_processos
            if _normalizar_texto(p.get("editor", "")) == nome_norm
            or _normalizar_texto(p.get("revisor", "")) == nome_norm
        ]

    sessoes_nao_distribuidas = {}
    sessoes_distribuidas = {}

    for p in todos_processos:
        numero_sessao = p.get("numero_sessao", "") or "Sem sessao"
        dia_sessao = _formatar_data_curta(p.get("dia_sessao"))
        tipo_sessao = p.get("tipo_sessao", "") or "Sem tipo"
        chave = f"{tipo_sessao} | Sessao {numero_sessao} | {dia_sessao}"

        if p.get("editor") or p.get("revisor"):
            if chave not in sessoes_distribuidas:
                sessoes_distribuidas[chave] = []
            sessoes_distribuidas[chave].append(p)
        else:
            if chave not in sessoes_nao_distribuidas:
                sessoes_nao_distribuidas[chave] = []
            sessoes_nao_distribuidas[chave].append(p)

    # --- SECAO 1: DISTRIBUIR ---
    st.markdown("### Distribuir Processos")

    if not sessoes_nao_distribuidas:
        st.info("Nao ha processos pendentes de distribuicao.")
    elif not modo_edicao:
        st.info("Modo visualizacao. A distribuicao pode ser executada apenas por gerentes ou superior.")
    else:
        disponiveis = _obter_disponiveis()

        if len(disponiveis) < 1:
            st.error(
                "Nao ha membros disponiveis para distribuir. "
                f"Atualmente ha {len(disponiveis)} membro(s) disponivel(is)."
            )
        else:
            chaves_distribuir = list(sessoes_nao_distribuidas.keys())
            sessao_sel = st.radio(
                "Selecione a sessao para distribuir",
                options=chaves_distribuir,
                key="dist_radio_distribuir",
            )

            processos_para_distribuir = sessoes_nao_distribuidas[sessao_sel]
            st.write(f"**{len(processos_para_distribuir)} processo(s)** para distribuir nesta sessao.")

            with st.expander("Ver processos", expanded=False):
                for p in processos_para_distribuir:
                    st.write(f"- {p.get('processo_numero', '')} | Relator: {p.get('relator', '-') or '-'}")

            st.markdown("### Selecionar Membros")
            st.caption("Desmarque os membros que nao devem participar desta distribuicao. No SEAT, todo mundo revisa todo mundo.")

            col_ed, col_rev = st.columns(2)

            editores_selecionados = []
            revisores_selecionados = []

            with col_ed:
                st.markdown("##### Editores")
                for membro in disponiveis:
                    if st.checkbox(membro, value=True, key=f"dist_ed_{membro}"):
                        editores_selecionados.append(membro)

            with col_rev:
                st.markdown("##### Revisores")
                for membro in disponiveis:
                    if st.checkbox(membro, value=True, key=f"dist_rev_{membro}"):
                        revisores_selecionados.append(membro)

            afastados = _obter_afastados()
            if afastados:
                with st.expander(f"{len(afastados)} membro(s) afastado(s) (excluido(s) automaticamente)"):
                    for nome in afastados:
                        st.write(f"- {nome}")

            st.markdown("---")
            if st.button("Distribuir", type="primary", use_container_width=True, key="dist_btn_distribuir"):
                if not editores_selecionados or not revisores_selecionados:
                    st.error("Selecione pelo menos 1 editor e 1 revisor.")
                else:
                    with st.spinner("Distribuindo processos..."):
                        atribuicoes = _distribuir_processos(
                            processos_para_distribuir,
                            editores_selecionados,
                            revisores_selecionados,
                        )

                    if atribuicoes:
                        salvos = 0
                        ignorados = 0
                        for proc_id, (editor, revisor) in atribuicoes.items():
                            # TRAVA OTIMISTA: Verifica se o processo ainda está sem atribuição
                            proc_atual = db_manager.buscar_por_id("pauta_seat", proc_id)
                            if proc_atual and (proc_atual.get("editor") or proc_atual.get("revisor")):
                                ignorados += 1
                                continue # Pula, pois outro gerente já distribuiu neste meio tempo

                            resultado = db_manager.atualizar("pauta_seat", proc_id, {
                                "editor": editor,
                                "revisor": revisor,
                            })
                            if resultado:
                                salvos += 1

                        if ignorados > 0:
                            st.warning(f"⚠️ Operação parcial: {ignorados} processo(s) já haviam sido distribuídos por outro colaborador no mesmo instante.")
                        if salvos > 0:
                            st.success(f"{salvos} processo(s) distribuído(s) com sucesso!")
                        st.rerun()
                    else:
                        st.warning("Nenhum processo foi distribuido. Verifique as condicoes.")

    st.markdown("---")

    # --- SECAO 2: EDITAR DISTRIBUICAO ---
    st.markdown("### Editar Distribuicao")

    if not sessoes_distribuidas:
        st.info("Nao ha processos distribuidos para editar.")
    else:
        chaves_editar = list(sessoes_distribuidas.keys())
        sessao_editar = st.radio(
            "Selecione a sessao para editar",
            options=chaves_editar,
            key="dist_radio_editar",
        )

        processos_distribuidos = sessoes_distribuidas[sessao_editar]
        disponiveis = _obter_equipe_seat()

        if not disponiveis:
            st.warning("Nao ha membros da equipe cadastrados.")
        else:
            if PANDAS_OK and modo_edicao:
                # Ordem correta: Processo - Relator - Editor - Editado - Revisor - Revisado - Comentario
                df_dados = []
                for p in processos_distribuidos:
                    df_dados.append({
                        "id": p.get("id"),
                        "processo_numero": p.get("processo_numero", ""),
                        "relator": p.get("relator", "") or "",
                        "editor": p.get("editor", "") or "",
                        "editado": bool(p.get("editado", False)),
                        "revisor": p.get("revisor", "") or "",
                        "revisado": bool(p.get("revisado", False)),
                        "comentario": p.get("comentario", "") or "",
                    })

                df = pd.DataFrame(df_dados)

                edited_df = st.data_editor(
                    df,
                    column_config={
                        "id": None,
                        "processo_numero": st.column_config.TextColumn("Processo", disabled=True),
                        "relator": st.column_config.TextColumn("Relator", disabled=True),
                        "editor": st.column_config.SelectboxColumn("Editor", options=disponiveis, required=True),
                        "editado": st.column_config.CheckboxColumn("Editado", default=False),
                        "revisor": st.column_config.SelectboxColumn("Revisor", options=disponiveis, required=True),
                        "revisado": st.column_config.CheckboxColumn("Revisado", default=False),
                        "comentario": st.column_config.TextColumn("Comentario", width="medium"),
                    },
                    hide_index=True,
                    use_container_width=True,
                    key="dist_data_editor_v2",
                )

                if st.button("Salvar Alteracoes", key="dist_btn_salvar"):
                    salvos = 0

                    for _, row in edited_df.iterrows():
                        editor_atual = row["editor"]
                        revisor_atual = row["revisor"]

                        original = next((p for p in processos_distribuidos if p.get("id") == row["id"]), None)
                        if not original:
                            continue

                        mudancas = {}

                        if (original.get("editor") or "") != editor_atual:
                            mudancas["editor"] = editor_atual if editor_atual else None

                        if (original.get("revisor") or "") != revisor_atual:
                            mudancas["revisor"] = revisor_atual if revisor_atual else None

                        editado_original = bool(original.get("editado", False))
                        editado_novo = bool(row["editado"])
                        if editado_original != editado_novo:
                            mudancas["editado"] = editado_novo

                        revisado_original = bool(original.get("revisado", False))
                        revisado_novo = bool(row["revisado"])
                        if revisado_original != revisado_novo:
                            mudancas["revisado"] = revisado_novo

                        comentario_original = original.get("comentario", "") or ""
                        comentario_novo = row["comentario"] or ""
                        if comentario_original != comentario_novo:
                            mudancas["comentario"] = comentario_novo.strip()

                        # Recalcular status automaticamente
                        if "editado" in mudancas or "revisado" in mudancas:
                            editado_val = mudancas.get("editado", editado_original)
                            revisado_val = mudancas.get("revisado", revisado_original)

                            if revisado_val and editado_val:
                                mudancas["status"] = "encaminhado"
                            elif editado_val and not revisado_val:
                                mudancas["status"] = "em_revisao"
                            elif not editado_val and not revisado_val:
                                mudancas["status"] = "em_edicao"
                            elif revisado_val and not editado_val:
                                mudancas["status"] = "encaminhado"

                        if mudancas:
                            resultado = db_manager.atualizar("pauta_seat", row["id"], mudancas)
                            if resultado:
                                salvos += 1

                    if salvos > 0:
                        st.success(f"{salvos} alteracao(oes) salva(s) com sucesso!")
                        st.rerun()
                    else:
                        st.info("Nenhuma alteracao detectada.")

            elif PANDAS_OK:
                # Modo visualizacao
                df_dados = []
                for p in processos_distribuidos:
                    df_dados.append({
                        "Processo": p.get("processo_numero", ""),
                        "Relator": p.get("relator", "") or "-",
                        "Editor": p.get("editor", "") or "-",
                        "Editado": "Sim" if p.get("editado", False) else "Nao",
                        "Revisor": p.get("revisor", "") or "-",
                        "Revisado": "Sim" if p.get("revisado", False) else "Nao",
                        "Comentario": p.get("comentario", "") or "-",
                    })
                df = pd.DataFrame(df_dados)
                st.dataframe(df, hide_index=True, use_container_width=True)

            else:
                for p in processos_distribuidos:
                    st.write(
                        f"- {p.get('processo_numero', '')} | "
                        f"Relator: {p.get('relator', '-') or '-'} | "
                        f"Editor: {p.get('editor', '-') or '-'} | "
                        f"Editado: {'Sim' if p.get('editado') else 'Nao'} | "
                        f"Revisor: {p.get('revisor', '-') or '-'} | "
                        f"Revisado: {'Sim' if p.get('revisado') else 'Nao'} | "
                        f"Comentario: {p.get('comentario', '') or '-'}"
                    )

# ============================================================
# DESPACHOS SINGULARES - FUNCOES
# ============================================================

def _verificar_despacho_singular_tab(numero_processo):
    """Verifica se um processo está cadastrado na tab de Despachos Singulares."""
    try:
        proc_norm = _normalizar_numero_processo(numero_processo)
        despachos = db_manager.buscar_todos("despachos_ds") or []

        if not despachos:
            return False, ""

        for d in despachos:
            d_num = _normalizar_numero_processo(d.get("processo_numero", ""))
            if d_num == proc_norm:
                tipo = d.get("tipo", "")
                if "sustentacao" in _normalizar_texto(tipo):
                    return True, "Sustentação Oral"
                return True, "Despacho Singular"
        return False, ""
    except Exception:
        return False, ""

def _verificar_despacho_singular(processo_pauta):
    """Verifica se um processo da pauta é Despacho Singular."""
    if not processo_pauta:
        return False
    # Verificar pelo tipo de sessão
    tipo = _normalizar_texto(str(processo_pauta.get("tipo_sessao", "")))
    if "despacho" in tipo and "singular" in tipo:
        return True
    # Verificar pelas observações
    obs = _normalizar_texto(str(processo_pauta.get("observacoes", "")))
    if "despacho singular" in obs:
        return True
    # Verificar pelo relator (se tiver "subst" ou "substituto")
    relator = _normalizar_texto(str(processo_pauta.get("relator", "")))
    if "despacho" in relator:
        return True
    return False

def _mover_despacho_singular_para_urgentes():
    """Move todos os processos de Despacho Singular da tab para a lista de urgentes."""
    try:
        despachos = db_manager.buscar_todos("despachos_ds") or []
        urgentes_existentes = db_manager.buscar_todos("processos_urgentes") or []
        nums_existentes = set()
        for u in urgentes_existentes:
            nums_existentes.add(_normalizar_numero_processo(u.get("processo_numero", "")))

        for d in despachos:
            d_num = _normalizar_numero_processo(d.get("processo_numero", ""))
            if d_num and d_num not in nums_existentes:
                tipo = d.get("tipo", "Despacho Singular")
                motivo = "Sustentação Oral" if "sustentacao" in _normalizar_texto(tipo) else "Despacho Singular"

                # Buscar dados da sessão na pauta
                tipo_sessao = ""
                sessao_num = ""
                dia_sessao = ""
                try:
                    pauta = db_manager.buscar_todos("pauta_seat") or []
                    for p in pauta:
                        p_num = _normalizar_numero_processo(p.get("processo_numero", ""))
                        if p_num == d_num:
                            tipo_sessao = p.get("tipo_sessao", "")
                            sessao_num = str(p.get("numero_sessao", ""))
                            dia_sessao = str(p.get("dia_sessao", ""))[:10]
                            break
                except Exception:
                    pass

                db_manager.inserir("processos_urgentes", {
                    "processo_numero": d_num,
                    "relator": d.get("relator", "N/I") or "N/I",
                    "motivo": motivo,
                    "tipo_sessao": tipo_sessao,
                    "sessao_numero": sessao_num,
                    "dia_sessao": dia_sessao,
                })
                nums_existentes.add(d_num)
    except Exception:
        pass

def _verificar_ds_pendente(processo_numero):
    """
    Verifica se o processo tem um DS pendente.
    Retorna o tipo ('Despacho Singular' ou 'Sustentacao Oral') ou None.
    """
    if not processo_numero:
        return None

    proc_higienizado = _higienizar_numero_processo(processo_numero)

    resultados = db_manager.buscar_todos(
        "despachos_ds",
        filtros={"processo_numero": proc_higienizado, "status": "pendente"}
    )

    if resultados:
        return resultados[0].get("tipo", "Despacho Singular")

    return None

def _marcar_ds_distribuido(processo_numero):
    """Marca o DS como distribuido apos o processo entrar na pauta."""
    if not processo_numero:
        return

    proc_higienizado = _higienizar_numero_processo(processo_numero)

    resultados = db_manager.buscar_todos(
        "despachos_ds",
        filtros={"processo_numero": proc_higienizado, "status": "pendente"}
    )

    for ds in resultados:
        db_manager.atualizar("despachos_ds", ds["id"], {"status": "distribuido"})

def _identificar_ds_apos_inclusao(processo_numero, pauta_id):
    """
    Funcao gatilho: apos incluir um processo na pauta_seat, verifica se
    existe um DS pendente para ele. Se sim, preenche o comentario e marca
    o DS como distribuido.
    """
    if not processo_numero or not pauta_id:
        return

    ds_tipo = _verificar_ds_pendente(processo_numero)

    if ds_tipo:
        comentario_ds = ds_tipo.upper()
        db_manager.atualizar("pauta_seat", pauta_id, {"comentario": comentario_ds})
        _marcar_ds_distribuido(processo_numero)

def _cadastrar_despacho(usuario):
    """Formulario para cadastrar novo Despacho Singular."""
    st.markdown("#### Cadastrar Novo Despacho")

    nome_usuario = usuario.get("nome", "Sistema")

    if "ds_oficios_temp" not in st.session_state:
        st.session_state["ds_oficios_temp"] = []

    col1, col2 = st.columns(2)

    with col1:
        processo_numero = st.text_input(
            "Numero do Processo *",
            placeholder="Ex: 00600-0007999/2022-63-e",
            key="ds_processo_numero"
        )
        tipo = st.selectbox(
            "Tipo *",
            options=["Despacho Singular", "Sustentacao Oral"],
            key="ds_tipo"
        )

    with col2:
        relator = st.text_input(
            "Relator (opcional)",
            placeholder="Ex: AM ou GAVF / Subst.",
            key="ds_relator"
        )
        forma_envio = st.selectbox(
            "Forma de Envio *",
            options=["E-mail", "Mensageria", "Protocolo"],
            key="ds_forma_envio"
        )

    observacoes = st.text_area(
        "Observacoes (opcional)",
        placeholder="Informacoes adicionais...",
        height=60,
        key="ds_observacoes"
    )

    st.markdown("---")
    st.markdown("##### Documentos Vinculados (Oficios / Memorandos)")
    st.caption("E obrigatorio cadastrar pelo menos 1 documento.")

    # Mostrar documentos ja adicionados
    if st.session_state["ds_oficios_temp"]:
        for i, of in enumerate(st.session_state["ds_oficios_temp"]):
            col_a, col_b, col_c, col_d, col_e = st.columns([2, 2, 2, 2, 1])
            with col_a:
                st.write(f"**{of['tipo_documento']}**")
            with col_b:
                st.write(f"N: {of['numero_oficio']}")
            with col_c:
                st.write(f"Para: {of['destinatario']}")
            with col_d:
                st.write(f"Envio: {of['tipo_envio']}")
            with col_e:
                if st.button("X", key=f"rm_of_{i}", help="Remover"):
                    st.session_state["ds_oficios_temp"].pop(i)
                    st.rerun()
    else:
        st.info("Nenhum documento adicionado ainda.")

    # Form para adicionar documento
    with st.form("form_add_oficio_ds"):
        st.markdown("**Adicionar Documento**")
        col_of1, col_of2 = st.columns(2)
        with col_of1:
            tipo_doc = st.selectbox(
                "Tipo de Documento",
                options=["Oficio", "Memorando"],
                key="ds_oficio_tipo"
            )
            numero_oficio = st.text_input(
                "Numero do Documento *",
                placeholder="Ex: 123/2026",
                key="ds_oficio_numero"
            )
        with col_of2:
            destinatario = st.text_input(
                "Destinatario *",
                placeholder="Ex: Secretaria de Fazenda",
                key="ds_oficio_dest"
            )
            tipo_envio_of = st.selectbox(
                "Tipo de Envio",
                options=["E-mail", "Mensageria", "Protocolo"],
                key="ds_oficio_envio"
            )

        adicionar = st.form_submit_button("Adicionar Documento")

        if adicionar:
            if not numero_oficio or not destinatario:
                st.error("Preencha o numero e o destinatario do documento.")
            else:
                st.session_state["ds_oficios_temp"].append({
                    "tipo_documento": tipo_doc,
                    "numero_oficio": numero_oficio,
                    "destinatario": destinatario,
                    "tipo_envio": tipo_envio_of,
                    "status": "aguardando",
                })
                st.rerun()

    st.markdown("---")

    # Botao para salvar o DS completo
    if st.button("Salvar Despacho", type="primary", use_container_width=True, key="btn_salvar_ds"):
        if not processo_numero:
            st.error("Numero do processo e obrigatorio.")
        elif not st.session_state["ds_oficios_temp"]:
            st.error("E obrigatorio cadastrar pelo menos 1 oficio/memorando.")
        else:
            proc_higienizado = _higienizar_numero_processo(processo_numero)
            relator_higienizado = _higienizar_relator(relator) if relator else None

            # Verificar se ja existe DS pendente
            ds_existente = _verificar_ds_pendente(proc_higienizado)
            if ds_existente:
                st.error(f"Ja existe um DS pendente para o processo {proc_higienizado}.")
                return

            dados_ds = {
                "processo_numero": proc_higienizado,
                "relator": relator_higienizado,
                "tipo": tipo,
                "forma_envio": forma_envio,
                "recebido_confirmado": False,
                "cadastrado_por": nome_usuario,
                "observacoes": observacoes.strip(),
                "status": "pendente",
            }

            resultado = db_manager.inserir("despachos_ds", dados_ds)

            if resultado:
                ds_id = resultado.get("id")

                if ds_id:
                    salvos = 0
                    for of in st.session_state["ds_oficios_temp"]:
                        dados_of = {
                            "despacho_id": ds_id,
                            "tipo_documento": of["tipo_documento"],
                            "numero_oficio": of["numero_oficio"],
                            "destinatario": of["destinatario"],
                            "tipo_envio": of["tipo_envio"],
                            "status": of["status"],
                        }
                        res_of = db_manager.inserir("oficios_ds", dados_of)
                        if res_of:
                            salvos += 1

                    st.session_state["ds_oficios_temp"] = []
                    st.success(f"Despacho cadastrado com sucesso! {salvos} documento(s) vinculado(s).")
                    st.rerun()
                else:
                    st.success("Despacho cadastrado, mas houve erro ao vincular documentos. Adicione manualmente na lista.")
                    st.rerun()
            else:
                st.error("Erro ao cadastrar despacho. Tente novamente.")

def _listar_despachos(usuario, modo_edicao):
    """Lista todos os DS cadastrados com seus documentos vinculados."""
    st.markdown("#### Despachos Cadastrados")

    nome_usuario = usuario.get("nome", "")

    todos_ds = db_manager.buscar_todos(
       "despachos_ds",
        ordem_coluna="created_at",
        ordem_desc=True,  # ← MUDAR DE False PARA True
       )

    col_f1, col_f2 = st.columns(2)
    with col_f1:
        filtro_status = st.selectbox(
            "Filtrar por status",
            options=["todos", "pendente", "distribuido"],
            key="ds_filtro_status"
        )
    with col_f2:
        busca = st.text_input(
            "Buscar por processo",
            placeholder="Digite o numero...",
            key="ds_busca"
        )

    if filtro_status != "todos":
        todos_ds = [d for d in todos_ds if d.get("status") == filtro_status]

    if busca.strip():
        busca_lower = busca.strip().lower()
        todos_ds = [d for d in todos_ds if busca_lower in (d.get("processo_numero", "") or "").lower()]

    if not todos_ds:
        st.info("Nenhum despacho cadastrado.")
        return

    st.write(f"**{len(todos_ds)} despacho(s) encontrado(s).**")

    for ds in todos_ds:
        status_icone = "🟡" if ds.get("status") == "pendente" else "🟢"
        # Exibir alterado_por se existir, senao cadastrado_por
        exibido_por = ds.get("alterado_por") or ds.get("cadastrado_por", "")
        rotulo_por = "Alterado por" if ds.get("alterado_por") else "Cadastrado por"

        with st.expander(
            f"{status_icone} {ds.get('processo_numero', '')} | "
            f"{ds.get('tipo', '')} | "
            f"{ds.get('status', '').upper()} | "
            f"{rotulo_por}: {exibido_por}"
        ):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.write(f"**Processo:** {ds.get('processo_numero', '')}")
                st.write(f"**Relator:** {ds.get('relator', '-') or '-'}")
            with col2:
                st.write(f"**Tipo:** {ds.get('tipo', '')}")
                st.write(f"**Envio:** {ds.get('forma_envio', '-') or '-'}")
            with col3:
                st.write(f"**{rotulo_por}:** {exibido_por}")
                st.write(f"**Recebido:** {'Sim' if ds.get('recebido_confirmado') else 'Nao'}")

            if ds.get("observacoes"):
                st.write(f"**Observacoes:** {ds.get('observacoes')}")

            oficios = db_manager.buscar_todos(
                "oficios_ds",
                filtros={"despacho_id": ds["id"]},
                ordem_coluna="created_at",
                ordem_desc=False,
            )

            st.markdown("**Documentos Vinculados:**")
            if oficios:
                for of in oficios:
                    col_a, col_b, col_c, col_d, col_e = st.columns([2, 2, 2, 2, 1])
                    with col_a:
                        st.write(f"{of.get('tipo_documento', '')} {of.get('numero_oficio', '')}")
                    with col_b:
                        st.write(f"Para: {of.get('destinatario', '')}")
                    with col_c:
                        st.write(f"Envio: {of.get('tipo_envio', '')}")
                    with col_d:
                        st.write(f"Status: {of.get('status', '')}")
                    with col_e:
                        if modo_edicao and of.get("status") == "aguardando":
                            if st.button("OK", key=f"of_recv_{of['id']}", help="Marcar como recebido"):
                                db_manager.atualizar("oficios_ds", of["id"], {"status": "recebido"})
                                db_manager.atualizar("despachos_ds", ds["id"], {"alterado_por": nome_usuario})
                                st.rerun()
            else:
                st.caption("Nenhum documento vinculado.")

            if modo_edicao:
                # Form para editar dados do DS
                with st.form(f"form_edit_ds_{ds['id']}"):
                    st.markdown("**Editar Despacho**")
                    col_e1, col_e2 = st.columns(2)
                    with col_e1:
                        edit_processo = st.text_input(
                            "Processo",
                            value=ds.get("processo_numero", ""),
                            key=f"edit_proc_{ds['id']}"
                        )
                        edit_relator = st.text_input(
                            "Relator",
                            value=ds.get("relator", "") or "",
                            key=f"edit_rel_{ds['id']}"
                        )
                        edit_forma = st.selectbox(
                            "Forma de Envio",
                            options=["E-mail", "Mensageria", "Protocolo"],
                            index=["E-mail", "Mensageria", "Protocolo"].index(ds.get("forma_envio", "E-mail")) if ds.get("forma_envio") in ["E-mail", "Mensageria", "Protocolo"] else 0,
                            key=f"edit_forma_{ds['id']}"
                        )
                    with col_e2:
                        edit_tipo = st.selectbox(
                            "Tipo",
                            options=["Despacho Singular", "Sustentacao Oral"],
                            index=["Despacho Singular", "Sustentacao Oral"].index(ds.get("tipo", "Despacho Singular")) if ds.get("tipo") in ["Despacho Singular", "Sustentacao Oral"] else 0,
                            key=f"edit_tipo_{ds['id']}"
                        )
                        edit_obs = st.text_area(
                            "Observacoes",
                            value=ds.get("observacoes", "") or "",
                            height=60,
                            key=f"edit_obs_{ds['id']}"
                        )

                    if st.form_submit_button("Salvar Alteracoes"):
                        proc_higienizado = _higienizar_numero_processo(edit_processo) if edit_processo else ds.get("processo_numero", "")
                        relator_higienizado = _higienizar_relator(edit_relator) if edit_relator else None
                        db_manager.atualizar("despachos_ds", ds["id"], {
                            "processo_numero": proc_higienizado,
                            "relator": relator_higienizado,
                            "tipo": edit_tipo,
                            "forma_envio": edit_forma,
                            "observacoes": edit_obs.strip(),
                            "alterado_por": nome_usuario,
                        })
                        st.success("Alteracoes salvas!")
                        st.rerun()

                # Form para adicionar documento extra
                with st.form(f"form_add_oficio_extra_{ds['id']}"):
                    st.markdown("**Adicionar Documento**")
                    col_a1, col_a2 = st.columns(2)
                    with col_a1:
                        novo_tipo_doc = st.selectbox(
                            "Tipo",
                            options=["Oficio", "Memorando"],
                            key=f"extra_tipo_{ds['id']}"
                        )
                        novo_numero = st.text_input(
                            "Numero *",
                            key=f"extra_numero_{ds['id']}"
                        )
                    with col_a2:
                        novo_dest = st.text_input(
                            "Destinatario *",
                            key=f"extra_dest_{ds['id']}"
                        )
                        novo_envio = st.selectbox(
                            "Tipo de Envio",
                            options=["E-mail", "Mensageria", "Protocolo"],
                            key=f"extra_envio_{ds['id']}"
                        )

                    if st.form_submit_button("Adicionar"):
                        if novo_numero and novo_dest:
                            db_manager.inserir("oficios_ds", {
                                "despacho_id": ds["id"],
                                "tipo_documento": novo_tipo_doc,
                                "numero_oficio": novo_numero,
                                "destinatario": novo_dest,
                                "tipo_envio": novo_envio,
                                "status": "aguardando",
                            })
                            db_manager.atualizar("despachos_ds", ds["id"], {"alterado_por": nome_usuario})
                            st.success("Documento adicionado!")
                            st.rerun()
                        else:
                            st.error("Preencha numero e destinatario.")

                col_btn1, col_btn2 = st.columns(2)
                with col_btn1:
                    if not ds.get("recebido_confirmado"):
                        if st.button("Marcar Recebido", key=f"ds_recv_{ds['id']}"):
                            db_manager.atualizar("despachos_ds", ds["id"], {
                                "recebido_confirmado": True,
                                "alterado_por": nome_usuario,
                            })
                            st.rerun()
                with col_btn2:
                    if ds.get("status") == "pendente":
                        if st.button("Marcar como Distribuido", key=f"ds_dist_{ds['id']}"):
                            db_manager.atualizar("despachos_ds", ds["id"], {
                                "status": "distribuido",
                                "alterado_por": nome_usuario,
                            })
                            st.rerun()

# ==================== MOTOR NIP INTELIGENTE ====================

def _extrair_texto_pdf(arquivo):
    """Extrai texto de um arquivo PDF usando pdfplumber."""
    try:
        import pdfplumber
        texto = ""
        with pdfplumber.open(arquivo) as pdf:
            for pagina in pdf.pages:
                pagina_texto = pagina.extract_text()
                if pagina_texto:
                    texto += pagina_texto + "\n"
        return texto
    except ImportError:
        st.error("Biblioteca 'pdfplumber' não instalada. Execute: pip install pdfplumber")
        return None
    except Exception as e:
        st.error(f"Erro ao extrair texto do PDF: {str(e)}")
        return None

def _identificar_relator(texto):
    """Identifica o relator pelo cabeçalho do PDF."""
    texto_upper = texto.upper()
    if "ANILCÉIA MACHADO" in texto_upper or "GABINETE DA CONSELHEIRA ANILCÉIA" in texto_upper:
        return "GCAM"
    elif "VINÍCIUS FRAGOSO" in texto_upper:
        return "VINICIUS_FRAGOSO"
    else:
        return "OUTRO"

def _extrair_voto(texto):
    """Extrai somente a parte do voto principal do texto completo do PDF."""
    import re

    # Padrões prioritários (voto principal — sempre têm "no sentido de que" + "egrégio")
    padroes_prioritarios = [
        r'(?:Diante do exposto|Pelo exposto|Ante o exposto)[,\s]*(?:em harmonia com o órgão instrutivo,?\s*)?VOTO\s*(?:no sentido de que\s*)?(?:por\s+o\s*)?(?:o\s+)?egrégio\s+(?:Tribunal|Plenário)[:\s]*',
        r'VOTO\s*(?:no sentido de que\s*)?(?:por\s+o\s*)?(?:o\s+)?egrégio\s+(?:Tribunal|Plenário)[:\s]*',
        r'(?:Diante do exposto|Pelo exposto|Ante o exposto)[,\s]*VOTO\s*(?:no sentido de que\s*)?(?:por\s+o\s*)?',
        r'VOTO\s*(?:no sentido de que\s*)?(?:por\s+o\s*)?',
    ]

    # Padrões de Voto de Vista que devem ser IGNORADOS
    padrao_voto_vista = re.compile(
        r'VOTO\s+em\s+harmonia',
        re.IGNORECASE
    )

    inicio_voto = None

    for padrao in padroes_prioritarios:
        matches = list(re.finditer(padrao, texto, re.IGNORECASE))
        if matches:
            # Filtrar: pegar o primeiro match que NÃO seja um Voto de Vista
            for match in matches:
                trecho_apos = texto[match.start():match.start() + 80]
                if not padrao_voto_vista.search(trecho_apos):
                    inicio_voto = match.start()
                    break
            if inicio_voto is not None:
                break

    if inicio_voto is None:
        return None

    voto = texto[inicio_voto:]

    # Remover assinatura no final
    assinatura_padroes = [
        r'\n\s*Sala das Sessões.*',
        r'\n\s*(?:Conselheir[oa]|Auditor|Conselheiro-Substituto)\s+',
        r'\n\s*[A-ZÀ-Ú]{4,}\s*\n\s*[A-ZÀ-Ú]',
    ]
    for ass_padrao in assinatura_padroes:
        ass_match = re.search(ass_padrao, voto, re.DOTALL)
        if ass_match:
            voto = voto[:ass_match.start()]

    return voto.strip()

def _aplicar_substituicoes(texto, regras):
    """
    Aplica substituições de frases, termos E verbos cadastrados no banco.
    Possui Curingas Inteligentes para aniquilar variações gramaticais.
    """
    import re
    
    # 1. CURINGAS INTELIGENTES (Wildcards Nativos do Motor NIP)
    # Substitui "da decisão que vier a ser proferida/prolatada/exarada/tomada" 
    # por "desta decisão", independente do verbo que o relator inventar no final.
    texto = re.sub(
        r'\bda\s+decisão\s+que\s+vier\s+a\s+ser\s+\w+\b', 
        'desta decisão', 
        texto, 
        flags=re.IGNORECASE
    )
    
    # 2. REGRAS DO BANCO DE DADOS (Com tolerância a quebras de PDF)
    if not regras:
        return texto
        
    for regra in regras:
        if regra.get("ativo", True):
            procurar = str(regra.get("procurar", "")).strip()
            substituir = str(regra.get("substituir_por", "")).strip()
            tipo = regra.get("tipo", "frase")
            
            if not procurar:
                continue

            # MÁGICA DE TOLERÂNCIA: Substitui espaços da regra por um padrão \s+ 
            # Isso faz com que a regra encontre a frase mesmo se o PDF quebrar a linha no meio.
            padrao_flexivel = r'\s+'.join([re.escape(p) for p in procurar.split()])
            
            # Se for termo isolado ou verbo, garante que pega a palavra inteira (\b)
            if tipo in ("termo", "verbo"):
                padrao_flexivel = rf'\b{padrao_flexivel}\b'
                
            texto = re.sub(padrao_flexivel, substituir, texto, flags=re.IGNORECASE)
            
    return texto

def _converter_imperativo_para_infinitivo_algoritmico(palavra, excecoes_banco=None):
    """
    FÓRMULA INTELIGENTE: Converte verbos no imperativo para o infinitivo usando regras morfossintáticas
    e radicais gramaticais da língua portuguesa.
    """
    p_lower = palavra.lower().strip()

    # 1. Checar exceções e verbos irregulares no banco ou mapa padrão
    irregulares = {
        "faça": "fazer", "faca": "fazer",
        "dê": "dar", "de": "dar",
        "seja": "ser",
        "veja": "ver",
        "vênia": "vênia", # Não alterar substantivos comuns
        "abstenha": "abster",
        "mantenha": "manter",
        "requer": "requerer",
    }
    
    if excecoes_banco and p_lower in excecoes_banco:
        return excecoes_banco[p_lower]
    if p_lower in irregulares:
        return irregulares[p_lower]

    # Se a palavra já termina com 'r' (já está no infinitivo), manter
    if p_lower.endswith('r'):
        return p_lower

    # 2. Remoção de Ífens e Pronomes Enclíticos (ex: "notifique-se" -> "notifique")
    pronomes = ["-se", "-nos", "-lhe", "-lhes", "-o", "-a", "-os", "-as"]
    for pronome in pronomes:
        if p_lower.endswith(pronome):
            p_lower = p_lower[:-len(pronome)]
            break

    # 3. Regras Algorítmicas de Sufixos da Língua Portuguesa
    
    # Verbos de 3ª Conjugação (-ir): defira -> deferir, indefira -> indeferir, exija -> exigir
    if p_lower.endswith("fira"):
        return p_lower[:-4] + "ferir"
    if p_lower.endswith("xija"):
        return p_lower[:-4] + "xigir"
    if p_lower.endswith("clua"):
        return p_lower[:-3] + "cluir"
    if p_lower.endswith("atenda"):
        return p_lower[:-4] + "tender"

    # Verbos de 1ª Conjugação (-ar) [A imensa maioria jurídica: determine, notifique, autorize, chame, tome]
    if p_lower.endswith("que"):  # comunique -> comunicar, notifique -> notificar
        return p_lower[:-3] + "car"
    if p_lower.endswith("gue"):  # homologue -> homologar, prorrogue -> prorrogar
        return p_lower[:-3] + "gar"
    if p_lower.endswith("ce"):   # autorize -> autorizar (variantes)
        return p_lower[:-2] + "çar"
    if p_lower.endswith("e"):    # determine -> determinar, tome -> tomar, considere -> considerar
        return p_lower[:-1] + "ar"
    if p_lower.endswith("em"):   # determinem -> determinar
        return p_lower[:-2] + "ar"

    # Verbos de 2ª Conjugação (-er)
    if p_lower.endswith("ceda"): # proceda -> proceder, conceda -> conceder
        return p_lower[:-4] + "ceder"
    if p_lower.endswith("mova"): # promova -> promover
        return p_lower[:-4] + "mover"
    if p_lower.endswith("a"):    # receba -> receber, responda -> responder
        return p_lower[:-1] + "er"

    return palavra # Caso não se enquadre, preserva o original

def _transformar_verbos(texto, regras):
    """
    Transforma verbos do Modo Imperativo Afirmativo para Infinitivo após numerais romanos e letras.
    Agora salta a palavra "não " para capturar o verbo corretamente.
    """
    import re

    excecoes_banco = {}
    if regras:
        for r in regras:
            if r.get("tipo") == "verbo" and r.get("ativo", True):
                excecoes_banco[r["procurar"].lower().strip()] = r["substituir_por"].lower().strip()

    # Regex flexível: aceita o "não " opcional antes do verbo e lida com vários tipos de traços/pontos
    padrao_romano = re.compile(
        r'((?:^|\n)\s*(?:[IVXLCDM]+|\d+)[\.\)\s]*[–\-—\.\•]*\s*(?:não\s+)?)(\w+(?:-\w+)?)',
        re.IGNORECASE
    )

    def _substituir_romano(match):
        prefixo = match.group(1)
        palavra = match.group(2)
        verbo_convertido = _converter_imperativo_para_infinitivo_algoritmico(palavra, excecoes_banco)
        
        if palavra[0].isupper():
            verbo_convertido = verbo_convertido.capitalize()
            
        return prefixo + verbo_convertido

    texto = padrao_romano.sub(_substituir_romano, texto)

    # Aplica o mesmo para letras de itens (a), b), etc)
    padrao_letra = re.compile(
        r'((?:^|\n)\s*[a-z]\)[\s]*(?:não\s+)?)(\w+(?:-\w+)?)',
        re.IGNORECASE
    )
    texto = padrao_letra.sub(_substituir_romano, texto)

    return texto

def _ofuscar_cpf(texto):
    """Ofusca CPFs no texto (3 primeiros e 2 últimos dígitos)."""
    import re
    def ofuscar(match):
        cpf = match.group(0)
        return f"***.{cpf[4:11]}-**"
    return re.sub(r'\d{3}\.\d{3}\.\d{3}-\d{2}', ofuscar, texto)

def _formatar_numerais(texto):
    """Padroniza formatação de numerais romanos, letras de itens e abreviações."""
    import re
    texto = re.sub(r'\b([IVXLCDM]+)[\.\)]\s*', r'\1 – ', texto)
    texto = re.sub(r'\b([IVXLCDM]+)\s*-\s*', r'\1 – ', texto)
    texto = re.sub(r'\b([a-z])\.\s*', r'\1) ', texto)
    texto = texto.replace("n.º", "nº").replace("n°", "nº")
    texto = re.sub(r'nº(\d)', r'nº \1', texto)
    texto = re.sub(r'LTDA(?!\.)\s*[-–]\s*ME', 'LTDA. – ME', texto)
    return texto

def _remover_e_antes_itens(texto):
    """Remove 'e' antes de itens (I, II, a), b))."""
    import re
    texto = re.sub(r';\s+e\s+([IVXLCDM]+)', r'; \1', texto)
    texto = re.sub(r';\s+e\s+([a-z])\)', r'; \1)', texto)
    return texto

def _limpar_cabecalho_rodape(texto):
    """Remove cabeçalhos e rodapés de PDFs de múltiplas páginas."""
    import re
    linhas = texto.split('\n')
    linhas_limpas = []
    for linha in linhas:
        linha_strip = linha.strip()
        linha_lower = linha_strip.lower()
        if 'documento assinado digitalmente' in linha_lower: continue
        if 'para verificar as assinaturas' in linha_lower: continue
        if 'acesse www.tc.df.gov.br' in linha_lower: continue
        if 'acesse www.tc.df.gov' in linha_lower: continue
        if linha_lower.startswith('e-doc'): continue
        if linha_lower.startswith('proc ') and '-' in linha_lower: continue
        if linha_lower == 'tribunal de contas do distrito federal': continue
        if linha_lower.startswith('gabinete da conselheir'): continue
        if linha_lower.startswith('gabinete do conselheir'): continue
        if linha_lower.startswith('gabinete do auditor'): continue
        if re.match(r'^e-?doc\s*\w+$', linha_strip, re.IGNORECASE): continue
        linhas_limpas.append(linha)

    texto = '\n'.join(linhas_limpas)
    return re.sub(r'\n{3,}', '\n\n', texto).strip()

def _adicionar_preambulo(texto, relator):
    """Adiciona o preâmbulo correto baseado no relator e remove o texto original do voto."""
    import re
    if relator == "GCAM":
        preambulo = "O Tribunal, por unanimidade, de acordo com o voto da Relatora, decidiu:"
    elif relator == "VINICIUS_FRAGOSO":
        preambulo = "O Tribunal, por unanimidade, de acordo com o voto do Relator, Conselheiro-Substituto VINÍCIUS FRAGOSO, atuando em substituição ao Conselheiro TAL, nos termos do art. 44, § 3º, do RI/TCDF, decidiu:"
    else:
        preambulo = "O Tribunal, por unanimidade, de acordo com o voto do Relator, decidiu:"

    padroes_remocao = [
        r'^\s*(?:Diante do exposto|Pelo exposto|Ante o exposto)[,\s]*(?:em harmonia com o órgão instrutivo,?\s*)?VOTO\s*(?:no sentido de que\s*)?(?:por\s+o\s*)?(?:o\s+)?egrégio\s+(?:Tribunal|Plenário)[:\s]*',
        r'^\s*VOTO\s*(?:no sentido de que\s*)?(?:por\s+o\s*)?(?:o\s+)?egrégio\s+(?:Tribunal|Plenário)[:\s]*',
        r'^\s*(?:Diante do exposto|Pelo exposto|Ante o exposto)[,\s]*VOTO\s*(?:no sentido de que\s*)?(?:por\s+o\s*)?',
        r'^\s*VOTO\s*(?:no sentido de que\s*)?(?:por\s+o\s*)?',
    ]
    for padrao in padroes_remocao:
        texto = re.sub(padrao, '', texto, flags=re.IGNORECASE)

    return f"{preambulo} {texto.strip()}"

def _corrigir_hifenizacao(texto):
    """Remove hifens que quebram palavras no final de linhas (comum em extração de PDF)."""
    import re
    return re.sub(r'(\w)-\s*\n\s*(\w)', r'\1\2', texto)
  
def _formatar_teleprompt(texto):
    """Formata o texto como teleprompt (texto corrido sem quebras)."""
    import re
    texto = re.sub(r'\n+', ' ', texto)
    texto = re.sub(r'\t+', ' ', texto)
    texto = re.sub(r'\r+', ' ', texto)
    texto = re.sub(r' +', ' ', texto)
    return texto.strip()

def _aplicar_negrito(texto):
    """Aplica negrito nas palavras-chave conforme as regras do NIP."""
    import re
    
    # 1. Numerais e itens
    texto = re.sub(r'\b([IVXLCDM]+)\s*–', r'**\1** –', texto)
    texto = re.sub(r'\b([a-z])\)', r'**\1)**', texto)
    
    for i in range(21):
        texto = re.sub(
            rf'\b{i}\s*(?:\([^)]+\)\s*)?dias?\b',
            lambda m: f'**{m.group(0)}**',
            texto,
            flags=re.IGNORECASE
        )

    palavras_negrito = [
        "urgente", "urgência", "prioritário", "prioridade", "brevidade",
        "imediato", "imediatamente", "importância",
        "suspender", "suspensão", "revoga", "abster", "abstenção",
        "anula", "anular", "negar",
        "continuidade", "continuação", "prosseguimento", "reabertura", "abertura",
        "licita", "licitação", "licitatório", "certame", "homologar", "adjudicar",
        "governador", "chefe do poder",
        "despacho singular", "sustentação oral",
        "audiência", "acórdão", "acórdãos", "notificação", "notificar",
        "cientificação", "cientificar", "convocação",
        "Covid", "Corona", "prorrog", "aprovar", "minuta", "pagamento",
        "tomar conhecimento", "considerar", "determinar", "chamar",
        "recomendar", "autorizar", "prazo de",
    ]

    for palavra in palavras_negrito:
        padrao = re.compile(re.escape(palavra), re.IGNORECASE)
        texto = padrao.sub(lambda m: f"**{m.group(0)}**", texto)

    return texto

def _processar_voto(voto_texto, relator, regras):
    """Pipeline completo de processamento do voto."""
    texto = _limpar_cabecalho_rodape(voto_texto)
    texto = _corrigir_hifenizacao(texto)
    
    # Aplica todas as regras do banco globalmente (incluindo verbos registrados manualmente)
    texto = _aplicar_substituicoes(texto, regras)
    
    # Executa o algoritmo inteligente para converter o resto dos verbos no início da linha
    texto = _transformar_verbos(texto, regras) 
    
    texto = _ofuscar_cpf(texto)
    texto = _formatar_numerais(texto)
    texto = _remover_e_antes_itens(texto)
    texto = _adicionar_preambulo(texto, relator)
    texto = _formatar_teleprompt(texto)
    texto = _aplicar_negrito(texto)

    return texto

def _gerar_docx(texto_markdown):
    """Gera um arquivo .docx a partir do texto com marcações markdown de negrito."""
    try:
        from docx import Document
        from io import BytesIO
        import re

        doc = Document()
        paragrafo = doc.add_paragraph()

        partes = re.split(r'(\*\*.*?\*\*)', texto_markdown)
        for parte in partes:
            if parte.startswith('**') and parte.endswith('**'):
                run = paragrafo.add_run(parte[2:-2])
                run.bold = True
            else:
                paragrafo.add_run(parte)

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    except ImportError:
        return None
    except Exception:
        return None

def _obter_regras_padrao():
    """Retorna regras padrão caso não existam no banco."""
    return [
        {"procurar": "Ministério Público junto ao Tribunal de Contas do Distrito Federal", "substituir_por": "Ministério Público junto ao Tribunal", "tipo": "frase", "ativo": True},
        {"procurar": "os acórdãos que submeto à apreciação plenária", "substituir_por": "os acórdãos apresentados pelo Relator", "tipo": "frase", "ativo": True},
        {"procurar": "o Acórdão ora submetido pelo Relator", "substituir_por": "os acórdãos apresentados pelo Relator", "tipo": "frase", "ativo": True},
        {"procurar": "da decisão que vier a ser proferida", "substituir_por": "desta decisão", "tipo": "frase", "ativo": True},
        {"procurar": "da decisão que vier a ser prolatada", "substituir_por": "desta decisão", "tipo": "frase", "ativo": True},
        {"procurar": "presente feito", "substituir_por": "feito em apreço", "tipo": "frase", "ativo": True},
        {"procurar": "presentes autos", "substituir_por": "autos em exame", "tipo": "frase", "ativo": True},
        {"procurar": "em comento", "substituir_por": "em análise", "tipo": "frase", "ativo": True},
        {"procurar": "Ministério Público de Contas", "substituir_por": "Ministério Público junto ao Tribunal - MPjTCDF", "tipo": "frase", "ativo": True},
        {"procurar": "n.º", "substituir_por": "nº", "tipo": "termo", "ativo": True},
        {"procurar": "n°", "substituir_por": "nº", "tipo": "termo", "ativo": True},
        {"procurar": "condutor", "substituir_por": "do Relator", "tipo": "frase", "ativo": True},
        {"procurar": "ciência da decisão que vier a ser proferida", "substituir_por": "ciência desta decisão", "tipo": "frase", "ativo": True},
        {"procurar": "ciência da decisão que vier a ser prolatada", "substituir_por": "ciência desta decisão", "tipo": "frase", "ativo": True},
    ]

# ==================== URGENTES E SERCON ====================

_RELATOR_SIGLA_MAP = {
    "MÁRCIO MICHEL": "GCMM",
    "MANOEL DE ANDRADE": "GCMA",
    "RENATO RAINHA": "GCRR",
    "ANILCÉIA MACHADO": "GCAM",
    "INÁCIO MAGALHÃES FILHO": "GCIM",
    "PAULO TADEU": "GCPT",
    "ANDRÉ CLEMENTE": "GCAC",
    "VINÍCIUS FRAGOSO": "GAVF",
}

def _normalizar_numero_processo(numero):
    """Normaliza o número do processo removendo sufixos, espaços e caracteres invisíveis."""
    if not numero:
        return ""
    numero = str(numero).strip()
    numero = numero.replace(" ", "").replace("\n", "").replace("\r", "").replace("\t", "")
    if numero.endswith("-e"): numero = numero[:-2]
    if numero.endswith("-E"): numero = numero[:-2]
    numero = numero.replace("\u200b", "").replace("\u00a0", "").replace("\ufeff", "")
    return numero.strip()

def _extrair_numero_processo(texto):
    """Extrai o numero do processo de forma altamente tolerante a formatos longos e curtos."""
    import re
    # Remove todos os espaços/quebras extras para formar uma string contínua
    texto_limpo = re.sub(r'\s+', ' ', texto)
    
    # 1. CAMADA: Tenta o padrão moderno/longo (Ex: 00600-00009313/2025-11-e)
    padrao_longo = r'(\d{4,5}\s*-\s*\d{6,8}\s*/\s*\d{4}\s*-\s*\d{2})(?:-e|-E)?'
    match_longo = re.search(padrao_longo, texto_limpo, re.IGNORECASE)
    if match_longo:
        return _normalizar_numero_processo(match_longo.group(1).replace(" ", ""))
    
    # 2. CAMADA: O Padrão Curto Ancorado (A sua sugestão!)
    # Caça variações: "Processo n.º: 25169/2017", "Proc.: 25169/2017", "Proc 25169/2017"
    padrao_ancora = r'(?:Processo|Proc\.?)\s*(?:n[º°.o]*\s*)?:?\s*(\d{2,8}\s*/\s*\d{4})(?:-e|-E)?'
    match_ancora = re.search(padrao_ancora, texto_limpo, re.IGNORECASE)
    if match_ancora:
        return _normalizar_numero_processo(match_ancora.group(1).replace(" ", ""))
        
    # 3. CAMADA DE SEGURANÇA: Procura qualquer número curto (XXXXX/YYYY) perdido no cabeçalho
    # Lê apenas os primeiros 1000 caracteres para não pegar leis ou artigos no meio do voto
    texto_inicio = texto_limpo[:1000]
    padrao_curto = r'\b(\d{2,8}\s*/\s*\d{4})(?:-e|-E)?\b'
    match_curto = re.search(padrao_curto, texto_inicio)
    if match_curto:
        return _normalizar_numero_processo(match_curto.group(1).replace(" ", ""))
        
    return None

def _identificar_relator_sigla(texto):
    """Identifica o relator pelo texto e retorna a sigla."""
    texto_upper = texto.upper()
    for nome, sigla in _RELATOR_SIGLA_MAP.items():
        if nome in texto_upper:
            return sigla
    return "N/I"

def _verificar_prazo(texto):
    """Verifica se ha prazos de 0 a 20 dias no texto."""
    import re
    padrao = r'\b(\d{1,2})\s*(?:\([^)]+\)\s*)?(?:dias?|dia)\b'
    matches = re.findall(padrao, texto, re.IGNORECASE)
    prazos = []
    for match in matches:
        try:
            num = int(match)
            if 0 <= num <= 20:
                prazos.append(f"{num} dias")
        except ValueError:
            pass
    return prazos

def _obter_palavras_urgencia():
    """Retorna a lista de palavras de urgencia do banco ou padrao."""
    try:
        palavras = db_manager.buscar_todos("palavras_urgencia_nip", filtros={"ativo": True})
        if palavras:
            return list(set([p["palavra"] for p in palavras]))
    except Exception:
        pass
    return [
        "urgente", "urgência", "prioritário", "prioridade", "brevidade",
        "importância", "imediato", "imediatamente", "suspender", "suspensão",
        "revoga", "abster", "abstenção", "anula", "anular", "negar",
        "continuidade", "continuação", "prosseguimento", "reabertura", "abertura",
        "licita", "licitação", "licitatório", "certame", "homologar", "adjudicar",
        "governador", "chefe do poder", "referendar", "ratificar",
        "despacho singular", "sustentação oral", "Covid", "Corona",
        "prorrog", "aprovar", "minuta", "pagamento", "prazo de",
    ]

def _obter_palavras_sercon():
    """Retorna a lista de palavras de SERCON do banco ou padrao."""
    try:
        palavras = db_manager.buscar_todos("palavras_sercon_nip", filtros={"ativo": True})
        if palavras:
            return palavras
    except Exception:
        pass
    return [
        {"palavra": "acórdão", "situacao": "acórdão"},
        {"palavra": "acórdãos", "situacao": "acórdão"},
        {"palavra": "notificação", "situacao": "notificação"},
        {"palavra": "notificar", "situacao": "notificação"},
        {"palavra": "cientificação", "situacao": "cientificação"},
        {"palavra": "cientificar", "situacao": "cientificação"},
        {"palavra": "audiência", "situacao": "audiência"},
        {"palavra": "convocação para audiência", "situacao": "audiência"},
    ]

def _verificar_urgencia(texto, palavras):
    """Verifica se o texto contem palavras de urgencia. Retorna (is_urgent, motivos)."""
    texto_lower = texto.lower()
    motivos = []

    for palavra in palavras:
        if palavra.lower() in texto_lower:
            motivos.append(palavra)

    # Verificar prazos de 0 a 20 dias
    prazos = _verificar_prazo(texto)
    motivos.extend(prazos)

    # Verificar Despacho Singular e Sustentação Oral no texto
    if "despacho singular" in texto_lower:
        if "Despacho Singular" not in motivos:
            motivos.append("Despacho Singular")
    if "sustentação oral" in texto_lower or "sustentacao oral" in texto_lower:
        if "Sustentação Oral" not in motivos:
            motivos.append("Sustentação Oral")

    # Deduplicar mantendo ordem
    motivos_unicos = []
    for m in motivos:
        if m not in motivos_unicos:
            motivos_unicos.append(m)

    return len(motivos_unicos) > 0, ", ".join(motivos_unicos) if motivos_unicos else ""

def _verificar_sercon(texto, palavras_sercon):
    """Verifica se o texto contem palavras de SERCON. Retorna (is_sercon, situacao)."""
    texto_lower = texto.lower()
    for item in palavras_sercon:
        palavra = item.get("palavra", "").lower()
        situacao = item.get("situacao", "")
        if palavra and palavra in texto_lower:
            return True, situacao
    return False, ""

def _renderizar_motor_nip(modo_edicao, usuario):
    """Funcao principal do Motor NIP - Edicao Automatica de Votos."""
    import streamlit as st
    st.markdown("### 🧠 Motor NIP Inteligente - Edição Automática de Votos")
    st.caption(
        "Faça upload do PDF do relatório/voto. O sistema extrai o voto, "
        "identifica os verbos gramaticalmente, aplica as regras de edição e entrega o texto pronto no formato teleprompt."
    )

    st.markdown("---")

    uploaded_file = st.file_uploader(
        "📄 Upload do PDF do Relatório/Voto",
        type=['pdf'],
        key="motor_nip_upload"
    )

    if uploaded_file is not None:
        with st.spinner("Processando PDF e analisando estrutura gramatical..."):
            texto_completo = _extrair_texto_pdf(uploaded_file)
            if not texto_completo:
                st.error("Não foi possível extrair texto do PDF.")
                return

            numero_processo = _extrair_numero_processo(texto_completo)
            relator_sigla = _identificar_relator_sigla(texto_completo)
            relator_tipo = _identificar_relator(texto_completo)

            relator_nome = {
                "GCAM": "GCAM (Anilcéia Machado — Relatora)",
                "VINICIUS_FRAGOSO": "GAVF (Vinícius Fragoso — Substituto)",
                "OUTRO": relator_sigla,
            }
            st.info(f"**Relator identificado:** {relator_nome.get(relator_tipo, relator_sigla)}")

            if numero_processo:
                st.info(f"**Processo identificado:** {numero_processo}")
            else:
                st.warning("Número do processo não encontrado automaticamente.")

            voto_extraido = _extrair_voto(texto_completo)
            if not voto_extraido:
                st.error("Não foi possível identificar o voto no PDF.")
                return

            try:
                regras = db_manager.buscar_todos(
                    "regras_substituicao_nip",
                    filtros={"ativo": True},
                    ordem_coluna="ordem",
                    ordem_desc=False,
                )
            except Exception:
                regras = []
            if not regras:
                regras = _obter_regras_padrao()

            texto_editado = _processar_voto(voto_extraido, relator_tipo, regras)
                       
            # Verificar urgencia e SERCON
            palavras_urg = _obter_palavras_urgencia()
            palavras_sercon = _obter_palavras_sercon()

            is_sercon, situacao_sercon = _verificar_sercon(voto_extraido, palavras_sercon)
            
            # Necessário invocar a checagem da aba de Despachos na SEAT
            from modulos.seat import _verificar_despacho_singular_tab
            is_ds, ds_motivo = _verificar_despacho_singular_tab(numero_processo or "")

            if is_sercon:
                is_urgent = False
                motivo_urg = ""
                st.warning(f"⚠️ Processo identificado para **SERCON** — Situação: {situacao_sercon}")
            elif is_ds:
                is_urgent = True
                motivo_urg = ds_motivo
                st.warning(f"⚠️ Processo identificado como **URGENTE** — Motivo: {ds_motivo}")
            else:
                is_urgent, motivo_urg = _verificar_urgencia(voto_extraido, palavras_urg)
                if is_urgent:
                    st.warning(f"⚠️ Processo identificado como **URGENTE** — Motivo: {motivo_urg}")

        st.markdown("---")
        st.markdown("#### ✅ Voto Editado")
        st.markdown(f"> {texto_editado}")

        st.markdown("---")

        texto_plain = texto_editado.replace("**", "")
        st.text_area(
            "📋 Texto para copiar (Ctrl+A, Ctrl+C):",
            value=texto_plain,
            height=300,
            key="motor_nip_resultado"
        )

        docx_data = _gerar_docx(texto_editado)
        if docx_data:
            st.download_button(
                label="📥 Baixar como Word (.docx) — com negrito",
                data=docx_data,
                file_name="voto_editado.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        st.markdown("---")
        st.markdown("#### 📋 Confirmação de Edição")

        col1, col2 = st.columns(2)
        with col1:
            proc_input = st.text_input(
                "Número do Processo",
                value=numero_processo or "",
                key="motor_nip_proc"
            )
        with col2:
            relator_input = st.text_input(
                "Relator (sigla)",
                value=relator_sigla,
                key="motor_nip_relator"
            )

        marcar_urgente = False
        if is_urgent and not is_sercon:
            marcar_urgente = st.checkbox(
                f"Marcar como urgente (motivo: {motivo_urg})",
                value=True,
                key="motor_nip_urgente"
            )

        if st.button("✅ Marcar como Editado", key="motor_nip_editado", type="primary"):
            if not proc_input:
                st.error("Informe o número do processo.")
                return

            proc_normalizado = _normalizar_numero_processo(proc_input)

            try:
                todos_pauta = db_manager.buscar_todos("pauta_seat") or []
                processo_pauta = None

                for p in todos_pauta:
                    p_num = _normalizar_numero_processo(p.get("processo_numero", ""))
                    if p_num == proc_normalizado:
                        processo_pauta = p
                        break

                if processo_pauta:
                    db_manager.atualizar("pauta_seat", processo_pauta["id"], {"editado": True})
                    st.success("✅ Processo marcado como editado na Pauta Ativa!")

                    if marcar_urgente:
                        db_manager.inserir("processos_urgentes", {
                            "processo_numero": proc_normalizado,
                            "relator": relator_input,
                            "motivo": motivo_urg,
                            "tipo_sessao": processo_pauta.get("tipo_sessao", ""),
                            "sessao_numero": str(processo_pauta.get("numero_sessao", "")),
                            "dia_sessao": str(processo_pauta.get("dia_sessao", ""))[:10],
                        })
                        st.success("✅ Processo adicionado à lista de Urgentes!")

                    if is_sercon:
                        db_manager.inserir("processos_sercon", {
                            "processo_numero": proc_normalizado,
                            "relator": relator_input,
                            "situacao": situacao_sercon,
                        })
                        st.success("✅ Processo enviado para o SERCON!")
                else:
                    nums_pauta = [_normalizar_numero_processo(p.get("processo_numero", "")) for p in todos_pauta]
                    st.warning(
                        f"Processo não encontrado na Pauta Ativa.\n\n"
                        f"**Número buscado:** `{proc_normalizado}`\n\n"
                        f"**Processos na pauta:** {', '.join(nums_pauta) if nums_pauta else 'NENHUM'}"
                    )
            except Exception as e:
                st.error(f"Erro ao atualizar: {str(e)}")

        with st.expander("Ver voto original extraído do PDF"):
            st.text(voto_extraido)

# ==================== URGENTES E SERCON ====================

_RELATOR_SIGLA_MAP = {
    "MÁRCIO MICHEL": "GCMM",
    "MANOEL DE ANDRADE": "GCMA",
    "RENATO RAINHA": "GCRR",
    "ANILCÉIA MACHADO": "GCAM",
    "INÁCIO MAGALHÃES FILHO": "GCIM",
    "PAULO TADEU": "GCPT",
    "ANDRÉ CLEMENTE": "GCAC",
    "VINÍCIUS FRAGOSO": "GAVF",
}

def _normalizar_numero_processo(numero):
    """Normaliza o número do processo removendo sufixos, espaços e caracteres invisíveis."""
    if not numero:
        return ""
    numero = str(numero).strip()
    numero = numero.replace(" ", "")
    numero = numero.replace("\n", "")
    numero = numero.replace("\r", "")
    numero = numero.replace("\t", "")
    if numero.endswith("-e"):
        numero = numero[:-2]
    if numero.endswith("-E"):
        numero = numero[:-2]
    numero = numero.replace("\u200b", "")
    numero = numero.replace("\u00a0", "")
    numero = numero.replace("\ufeff", "")
    return numero.strip()

def _extrair_numero_processo(texto):
    """Extrai o numero do processo do cabecalho do PDF."""
    import re
    # Padrao: Processo nº: 00600-00009313/2025-11-e (ou sem -e)
    padrao = r'Processo\s*n[º°.o]*\s*:?\s*\n?\s*(\d{5}-\d{8}/\d{4}-\d{2})(?:-e)?'
    match = re.search(padrao, texto, re.IGNORECASE)
    if match:
        return _normalizar_numero_processo(match.group(1))
    # Fallback: procurar o padrao do numero em qualquer lugar
    padrao_fallback = r'(\d{5}-\d{8}/\d{4}-\d{2})(?:-e)?'
    match = re.search(padrao_fallback, texto)
    if match:
        return _normalizar_numero_processo(match.group(1))
    return None

def _identificar_relator_sigla(texto):
    """Identifica o relator pelo texto e retorna a sigla."""
    texto_upper = texto.upper()
    for nome, sigla in _RELATOR_SIGLA_MAP.items():
        if nome in texto_upper:
            return sigla
    return "N/I"

def _verificar_prazo(texto):
    """Verifica se ha prazos de 0 a 20 dias no texto."""
    import re
    padrao = r'\b(\d{1,2})\s*(?:\([^)]+\)\s*)?(?:dias?|dia)\b'
    matches = re.findall(padrao, texto, re.IGNORECASE)
    prazos = []
    for match in matches:
        try:
            num = int(match)
            if 0 <= num <= 20:
                prazos.append(f"{num} dias")
        except ValueError:
            pass
    return prazos

def _obter_palavras_urgencia():
    """Retorna a lista de palavras de urgencia do banco ou padrao."""
    try:
        palavras = db_manager.buscar_todos("palavras_urgencia_nip", filtros={"ativo": True})
        if palavras:
            return list(set([p["palavra"] for p in palavras]))
    except Exception:
        pass
    return [
        "urgente", "urgência", "prioritário", "prioridade", "brevidade",
        "importância", "imediato", "imediatamente", "suspender", "suspensão",
        "revoga", "abster", "abstenção", "anula", "anular", "negar",
        "continuidade", "continuação", "prosseguimento", "reabertura", "abertura",
        "licita", "licitação", "licitatório", "certame", "homologar", "adjudicar",
        "governador", "chefe do poder", "referendar", "ratificar",
        "despacho singular", "sustentação oral", "Covid", "Corona",
        "prorrog", "aprovar", "minuta", "pagamento", "prazo de",
    ]

def _obter_palavras_sercon():
    """Retorna a lista de palavras de SERCON do banco ou padrao."""
    try:
        palavras = db_manager.buscar_todos("palavras_sercon_nip", filtros={"ativo": True})
        if palavras:
            return palavras
    except Exception:
        pass
    return [
        {"palavra": "acórdão", "situacao": "acórdão"},
        {"palavra": "acórdãos", "situacao": "acórdão"},
        {"palavra": "notificação", "situacao": "notificação"},
        {"palavra": "notificar", "situacao": "notificação"},
        {"palavra": "cientificação", "situacao": "cientificação"},
        {"palavra": "cientificar", "situacao": "cientificação"},
        {"palavra": "audiência", "situacao": "audiência"},
        {"palavra": "convocação para audiência", "situacao": "audiência"},
    ]

def _verificar_urgencia(texto, palavras):
    """Verifica se o texto contem palavras de urgencia. Retorna (is_urgent, motivos)."""
    texto_lower = texto.lower()
    motivos = []

    for palavra in palavras:
        if palavra.lower() in texto_lower:
            motivos.append(palavra)

    # Verificar prazos de 0 a 20 dias
    prazos = _verificar_prazo(texto)
    motivos.extend(prazos)

    # Verificar Despacho Singular e Sustentação Oral no texto
    if "despacho singular" in texto_lower:
        if "Despacho Singular" not in motivos:
            motivos.append("Despacho Singular")
    if "sustentação oral" in texto_lower or "sustentacao oral" in texto_lower:
        if "Sustentação Oral" not in motivos:
            motivos.append("Sustentação Oral")

    # Deduplicar mantendo ordem
    motivos_unicos = []
    for m in motivos:
        if m not in motivos_unicos:
            motivos_unicos.append(m)

    return len(motivos_unicos) > 0, ", ".join(motivos_unicos) if motivos_unicos else ""

def _verificar_sercon(texto, palavras_sercon):
    """Verifica se o texto contem palavras de SERCON. Retorna (is_sercon, situacao)."""
    texto_lower = texto.lower()
    for item in palavras_sercon:
        palavra = item.get("palavra", "").lower()
        situacao = item.get("situacao", "")
        if palavra and palavra in texto_lower:
            return True, situacao
    return False, ""

def _renderizar_urgentes(modo_edicao, usuario):
    """Renderiza a tab de Urgentes com 4 tabelas por tipo de sessao."""
    st.markdown("### Processos Urgentes")
    st.caption("Processos identificados como urgentes, organizados por tipo de sessao.")

    # Buscar processos urgentes
    try:
        urgentes = db_manager.buscar_todos("processos_urgentes")
    except Exception:
        urgentes = []

    if not urgentes:
        st.info("Nenhum processo urgente cadastrado ainda.")
        return

    # Buscar datas das ultimas sessoes de cada tipo
    todos_processos = db_manager.buscar_todos("pauta_seat") or []
    datas_recentes = {}
    for p in todos_processos:
        tipo = p.get("tipo_sessao", "")
        dia = p.get("dia_sessao")
        if not tipo or not dia:
            continue
        if "urgente" in _normalizar_texto(tipo):
            continue
        dia_str = str(dia)[:10]
        tipo_key = _normalizar_texto(tipo)
        if tipo_key not in datas_recentes or dia_str > datas_recentes[tipo_key]:
            datas_recentes[tipo_key] = dia_str

    # Filtrar urgentes apenas das ultimas sessoes
    urgentes_filtrados = []
    for u in urgentes:
        dia_u = str(u.get("dia_sessao", ""))[:10]
        tipo_u = _normalizar_texto(u.get("tipo_sessao", ""))
        if tipo_u in datas_recentes and dia_u == datas_recentes[tipo_u]:
            urgentes_filtrados.append(u)

    # Agrupar por tipo de sessao
    tipos_sessao = {
        "sessao ordinaria": "Sessão Ordinária",
        "sessao reservada": "Sessão Reservada",
        "sessao administrativa": "Sessão Administrativa",
        "sessao ordinaria virtual": "Sessão Ordinária Virtual",
    }

    for tipo_key, tipo_label in tipos_sessao.items():
        processos_tipo = [
            u for u in urgentes_filtrados
            if _normalizar_texto(u.get("tipo_sessao", "")) == tipo_key
        ]

        st.markdown(f"#### {tipo_label}")
        if processos_tipo:
            import pandas as pd
            dados = []
            for u in processos_tipo:
                dados.append({
                    "Processo": u.get("processo_numero", ""),
                    "Relator": u.get("relator", ""),
                    "Motivo": u.get("motivo", ""),
                })
            df = pd.DataFrame(dados)
            st.dataframe(df, hide_index=True, use_container_width=True, height=len(df) * 35 + 40)
        else:
            st.caption("Nenhum processo urgente nesta sessão.")
        st.markdown("---")

def _renderizar_sidebar_urgentes(usuario):
    """Mostra a tabela de urgentes na sidebar, abaixo dos Despachos Singulares."""
    import pandas as pd

    cargo = usuario.get("cargo", "operacional")
    if cargo not in ("criador", "raiz", "gerente"):
        return

    try:
        urgentes = db_manager.buscar_todos("processos_urgentes")
    except Exception:
        urgentes = []

    if not urgentes:
        return

    # Filtrar apenas das ultimas sessoes
    todos_processos = db_manager.buscar_todos("pauta_seat") or []
    datas_recentes = {}
    for p in todos_processos:
        tipo = p.get("tipo_sessao", "")
        dia = p.get("dia_sessao")
        if not tipo or not dia:
            continue
        if "urgente" in _normalizar_texto(tipo):
            continue
        dia_str = str(dia)[:10]
        tipo_key = _normalizar_texto(tipo)
        if tipo_key not in datas_recentes or dia_str > datas_recentes[tipo_key]:
            datas_recentes[tipo_key] = dia_str

    urgentes_filtrados = []
    for u in urgentes:
        dia_u = str(u.get("dia_sessao", ""))[:10]
        tipo_u = _normalizar_texto(u.get("tipo_sessao", ""))
        if tipo_u in datas_recentes and dia_u == datas_recentes[tipo_u]:
            urgentes_filtrados.append(u)

    if not urgentes_filtrados:
        return

    with st.sidebar:
        st.markdown("---")
        st.markdown("##### Urgentes (Recentes)")
        dados = []
        for u in urgentes_filtrados:
            dados.append({
                "Processo": u.get("processo_numero", ""),
                "Relator": u.get("relator", ""),
                "Motivo": u.get("motivo", ""),
            })
        df = pd.DataFrame(dados)
        st.dataframe(
            df,
            hide_index=True,
            use_container_width=True,
            height=len(df) * 35 + 40,
        )

# ==================== ESCALA DOE ====================

def _obter_config_doe():
    """Obtém a data de início da rotação do DOE."""
    from datetime import date, timedelta
    try:
        config = db_manager.buscar_todos("config_doe") or []
        if config:
            data_str = config[0].get("data_inicio_rotacao")
            if data_str:
                return date.fromisoformat(str(data_str)[:10])
    except Exception:
        pass
    # Default: segunda-feira da semana atual
    hoje = date.today()
    return hoje - timedelta(days=hoje.weekday())

def _salvar_config_doe(data_inicio):
    """Salva a data de início da rotação."""
    try:
        config = db_manager.buscar_todos("config_doe") or []
        if config:
            db_manager.atualizar("config_doe", config[0]["id"], {
                "data_inicio_rotacao": str(data_inicio),
            })
        else:
            db_manager.inserir("config_doe", {
                "data_inicio_rotacao": str(data_inicio),
            })
    except Exception:
        pass

def _obter_duplas_doe():
    """Retorna todas as duplas ativas ordenadas."""
    try:
        return db_manager.buscar_todos(
            "duplas_doe",
            filtros={"ativo": True},
            ordem_coluna="ordem",
            ordem_desc=False,
        ) or []
    except Exception:
        return []

def _obter_semana_atual():
    """Retorna (segunda, sexta) da semana atual."""
    from datetime import date, timedelta
    hoje = date.today()
    segunda = hoje - timedelta(days=hoje.weekday())
    sexta = segunda + timedelta(days=4)
    return segunda, sexta

def _obter_dupla_semana(data_referencia, duplas):
    """Retorna qual dupla está escalada para a semana da data_referencia."""
    from datetime import date, timedelta
    if not duplas:
        return None
    data_inicio = _obter_config_doe()
    dias_desde_segunda = data_inicio.weekday()
    segunda_inicio = data_inicio - timedelta(days=dias_desde_segunda)
    dias_diff = (data_referencia - segunda_inicio).days
    if dias_diff < 0:
        return None
    semanas_diff = dias_diff // 7
    indice = semanas_diff % len(duplas)
    return duplas[indice]

def _gerar_calendario_doe(duplas, num_semanas=8):
    """Gera o calendário de DOE para as próximas N semanas."""
    from datetime import date, timedelta
    calendario = []
    segunda_atual, _ = _obter_semana_atual()
    for i in range(num_semanas):
        segunda = segunda_atual + timedelta(weeks=i)
        sexta = segunda + timedelta(days=4)
        dupla = _obter_dupla_semana(segunda, duplas)
        calendario.append({
            "semana_inicio": segunda,
            "semana_fim": sexta,
            "dupla": dupla,
        })
    return calendario

def _verificar_conflitos_ferias():
    """Verifica todos os conflitos entre férias aprovadas e escala DOE."""
    from datetime import date, timedelta
    conflitos = []
    try:
        duplas = _obter_duplas_doe()
        if not duplas:
            return conflitos

        ferias_todas = db_manager.buscar_todos(
            "ferias_colaboradores",
            filtros={"status": "aprovada"},
        ) or []

        segunda_atual, _ = _obter_semana_atual()

        for ferias in ferias_todas:
            colaborador = ferias.get("colaborador", "")
            data_ini = date.fromisoformat(str(ferias.get("data_inicio"))[:10])
            data_fim = date.fromisoformat(str(ferias.get("data_fim"))[:10])

            for i in range(12):
                segunda = segunda_atual + timedelta(weeks=i)
                sexta = segunda + timedelta(days=4)

                if segunda <= data_fim and sexta >= data_ini:
                    dupla = _obter_dupla_semana(segunda, duplas)
                    if dupla and colaborador in [dupla.get("membro1"), dupla.get("membro2")]:
                        substituto = ferias.get("substituto", "")
                        conflitos.append({
                            "colaborador": colaborador,
                            "semana_inicio": segunda,
                            "semana_fim": sexta,
                            "dupla": dupla,
                            "ferias_inicio": data_ini,
                            "ferias_fim": data_fim,
                            "substituto": substituto,
                            "ferias_id": ferias.get("id"),
                        })
    except Exception:
        pass
    return conflitos

def _renderizar_sidebar_doe(usuario):
    """Mostra a dupla escalada na semana atual na barra lateral."""
    from datetime import date, timedelta
    cargo = usuario.get("cargo", "operacional")
    setor = usuario.get("setor", "")

    if cargo not in ("criador", "raiz", "gerente") and "seat" not in _normalizar_texto(setor):
        return

    duplas = _obter_duplas_doe()
    if not duplas:
        return

    segunda, sexta = _obter_semana_atual()
    dupla_atual = _obter_dupla_semana(segunda, duplas)
    if not dupla_atual:
        return

    with st.sidebar:
        st.markdown("---")
        st.markdown("##### 📰 Escala DOE (Esta Semana)")
        st.write(f"**Período:** {segunda.strftime('%d/%m')} a {sexta.strftime('%d/%m')}")

        # Verificar substitutos
        membro1 = dupla_atual.get("membro1", "")
        membro2 = dupla_atual.get("membro2", "")

        conflitos = _verificar_conflitos_ferias()
        conflitos_semana = [c for c in conflitos if c["semana_inicio"] == segunda]

        for c in conflitos_semana:
            if c["colaborador"] == membro1 and c.get("substituto"):
                membro1 = f"~~{membro1}~~ → {c['substituto']}"
            elif c["colaborador"] == membro2 and c.get("substituto"):
                membro2 = f"~~{membro2}~~ → {c['substituto']}"

        st.write(f"**Dupla:** {dupla_atual.get('nome_dupla', 'N/I')}")
        st.write(f"• {membro1}")
        st.write(f"• {membro2}")

        # Alertas de férias sem substituto
        for c in conflitos_semana:
            if not c.get("substituto"):
                st.warning(
                    f"⚠️ **{c['colaborador']}** está de férias "
                    f"({c['ferias_inicio'].strftime('%d/%m')} a {c['ferias_fim'].strftime('%d/%m')}). "
                    f"Sem substituto designado!"
                )

def _renderizar_escala_doe(modo_edicao, usuario):
    """Renderiza a aba de Escala DOE em modo visualização limpa."""
    from datetime import date, timedelta
    import pandas as pd

    st.markdown("### 📰 Escala DOE — Diário Oficial do Tribunal")
    st.caption(
        "Visualização do rodízio de duplas para a publicação do DOE. "
        "A gestão de duplas, regras de rodízio e atribuição de substitutos são administradas pelo Gabinete."
    )

    # === Seção 1: Semana atual ===
    st.markdown("---")
    st.markdown("#### 📌 Semana Atual")

    duplas = _obter_duplas_doe()
    segunda, sexta = _obter_semana_atual()

    if not duplas:
        st.info("Nenhuma dupla ativa no momento. A escala será disponibilizada após configuração pelo Gabinete.")
    else:
        dupla_atual = _obter_dupla_semana(segunda, duplas)
        if dupla_atual:
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Período", f"{segunda.strftime('%d/%m/%Y')} a {sexta.strftime('%d/%m/%Y')}")
            with col2:
                st.metric("Dupla Escalada", dupla_atual.get("nome_dupla", "N/I"))
            with col3:
                st.metric("Membros Titulares", f"{dupla_atual.get('membro1', '')} / {dupla_atual.get('membro2', '')}")

            # Verificar ausências (férias aprovadas ou atestados) na semana atual
            conflitos = _verificar_conflitos_ferias()
            conflitos_semana = [c for c in conflitos if c["semana_inicio"] == segunda]
            for c in conflitos_semana:
                if c.get("substituto"):
                    st.info(
                        f"🔄 **Substituição Ativa:** {c['colaborador']} está ausente no período. "
                        f"Substituto em atuação: **{c['substituto']}**."
                    )
                else:
                    st.warning(
                        f"⚠️ **Alerta Operacional:** {c['colaborador']} está ausente de "
                        f"{c['ferias_inicio'].strftime('%d/%m/%Y')} a {c['ferias_fim'].strftime('%d/%m/%Y')}. "
                        f"**Aguardando designação de substituto pelo Gabinete.**"
                    )
        else:
            st.warning("Não foi possível determinar a dupla da semana atual.")

    # === Seção 2: Próximas semanas ===
    st.markdown("---")
    st.markdown("#### 📅 Próximas Semanas")

    if duplas:
        calendario = _gerar_calendario_doe(duplas, num_semanas=8)
        dados_cal = []
        for item in calendario:
            dupla = item["dupla"]
            dados_cal.append({
                "Semana": f"{item['semana_inicio'].strftime('%d/%m')} a {item['semana_fim'].strftime('%d/%m')}",
                "Dupla": dupla.get("nome_dupla", "N/I") if dupla else "—",
                "Membro 1": dupla.get("membro1", "") if dupla else "—",
                "Membro 2": dupla.get("membro2", "") if dupla else "—",
            })
        df_cal = pd.DataFrame(dados_cal)
        st.dataframe(df_cal, hide_index=True, use_container_width=True)

# ============================================================
# SUBMÓDULO: FÉRIAS E AFASTAMENTOS (SEAT)
# ============================================================

def _verificar_radar_choques(data_ini, data_fim, id_ignorar=None):
    """
    Radar de Choques: verifica se há outros colaboradores da SEAT ausentes no mesmo período.
    Retorna lista de colisões encontradas.
    """
    from datetime import date
    try:
        solicitacoes = db_manager.buscar_todos("solicitacoes_ausencia") or []
    except Exception:
        return []

    choques = []
    for s in solicitacoes:
        if id_ignorar and str(s.get("id")) == str(id_ignorar):
            continue
        # Considera apenas pedidos aprovados ou atestados notificados
        if s.get("status") not in ("APROVADA", "NOTIFICADO"):
            continue

        s_ini = date.fromisoformat(str(s.get("data_inicio"))[:10])
        s_fim = date.fromisoformat(str(s.get("data_fim"))[:10])

        # Verifica intersecção de datas
        if data_ini <= s_fim and data_fim >= s_ini:
            choques.append({
                "colaborador": s.get("colaborador_nome", "Colaborador"),
                "tipo": s.get("tipo", "AUSENCIA"),
                "inicio": s_ini.strftime("%d/%m/%Y"),
                "fim": s_fim.strftime("%d/%m/%Y"),
            })
    return choques

def _verificar_radar_choques_seat(data_ini, data_fim, id_ignorar=None):
    """
    Radar de Choques: verifica se há outros colaboradores da SEAT ausentes no mesmo período.
    Retorna lista de colisões encontradas.
    """
    from datetime import date
    try:
        solicitacoes = db_manager.buscar_todos("solicitacoes_ausencia") or []
    except Exception:
        return []
    choques = []
    for s in solicitacoes:
        if id_ignorar and str(s.get("id")) == str(id_ignorar):
            continue
        if s.get("status") not in ("APROVADA", "NOTIFICADO"):
            continue
        if s.get("setor", "").upper() != "SEAT":
            continue
        s_ini = date.fromisoformat(str(s.get("data_inicio"))[:10])
        s_fim = date.fromisoformat(str(s.get("data_fim"))[:10])
        if data_ini <= s_fim and data_fim >= s_ini:
            tipo_label = "Férias" if s.get("tipo") == "FERIAS" else ("Atestado" if s.get("tipo") == "ATESTADO" else "Abono")
            choques.append({
                "colaborador": s.get("colaborador_nome", "Colaborador"),
                "tipo": tipo_label,
                "inicio": s_ini.strftime("%d/%m/%Y"),
                "fim": s_fim.strftime("%d/%m/%Y"),
            })
    return choques

def _renderizar_ausencias_seat(modo_edicao: bool, usuario: dict):
    """Renderiza a aba de Férias, Atestados e Abono da SEAT."""
    from datetime import date
    import pandas as pd

    # Proteção contra usuario None
    if not usuario or not isinstance(usuario, dict):
        st.warning("Não foi possível carregar os dados do usuário logado.")
        return

    st.markdown("### 🌴 Férias, Atestados e Abono")
    st.caption(
        "Solicitação de férias, registro de atestados médicos e pedido de abono. "
        "Férias e abonos são enviados para análise da chefia no Gabinete. "
        "Atestados médicos são notificados automaticamente."
    )

    # Identificação automática pelo login
    nome_usuario = usuario.get("nome", "Colaborador")
    matricula_usuario = str(usuario.get("matricula", ""))

    tab_solicitar, tab_quadro = st.tabs([
        "➕ Nova Solicitação",
        "📅 Quadro Público de Ausências",
    ])

    # --- ABA 1: NOVA SOLICITAÇÃO ---
    with tab_solicitar:
        st.markdown(f"**Colaborador Solicitante:** `{nome_usuario}` (Matrícula: `{matricula_usuario}`)")
        st.info("O sistema identifica seu perfil automaticamente. Selecione o tipo de registro abaixo.")

        tipo_registro = st.radio(
            "Tipo de Registro",
            ["Férias", "Atestado Médico", "Abono"],
            horizontal=True,
            key="tipo_registro_seat"
        )

        with st.form("form_registro_ausencia_seat"):
            col1, col2 = st.columns(2)
            with col1:
                data_ini = st.date_input("Data de Início *", value=date.today(), key="dt_ini_seat")
            with col2:
                data_fim = st.date_input("Data de Retorno / Fim *", value=date.today(), key="dt_fim_seat")

            observacoes = st.text_area(
                "Observações / Motivo",
                placeholder="Informações adicionais para a chefia ou equipe...",
                height=70,
                key="obs_ausencia_seat"
            )

            submit_ausencia = st.form_submit_button("Registrar no Sistema", type="primary", use_container_width=True)

            if submit_ausencia:
                if data_fim < data_ini:
                    st.error("A data de término não pode ser anterior à data de início.")
                else:
                    dias_total = (data_fim - data_ini).days + 1

                    # Definir tipo e status conforme o registro
                    if tipo_registro == "Férias":
                        tipo_db = "FERIAS"
                        status_inicial = "PENDENTE"
                    elif tipo_registro == "Atestado Médico":
                        tipo_db = "ATESTADO"
                        status_inicial = "NOTIFICADO"
                    else:  # Abono
                        tipo_db = "ABONO"
                        status_inicial = "PENDENTE"

                    # Radar de choques — verifica sobreposição com outros colaboradores
                    choques = _verificar_radar_choques_seat(data_ini, data_fim)

                    dados_ausencia = {
                        "matricula": matricula_usuario,
                        "colaborador_nome": nome_usuario,
                        "setor": "SEAT",
                        "tipo": tipo_db,
                        "data_inicio": data_ini.isoformat(),
                        "data_fim": data_fim.isoformat(),
                        "dias_afastado": dias_total,
                        "observacoes": observacoes.strip(),
                        "status": status_inicial,
                    }

                    res = db_manager.inserir("solicitacoes_ausencia", dados_ausencia)

                    if res:
                        if tipo_db == "FERIAS":
                            msg = f"✅ Solicitação de férias ({dias_total} dias) enviada para análise da chefia no Gabinete."
                        elif tipo_db == "ATESTADO":
                            msg = f"✅ Atestado médico ({dias_total} dias) notificado com sucesso e publicado no quadro!"
                        else:
                            msg = f"✅ Pedido de abono ({dias_total} dias) enviado para análise da chefia no Gabinete."
                        st.success(msg)

                        # Alertar sobre choques detectados
                        if choques:
                            st.warning(f"⚠️ **Atenção:** {len(choques)} colaborador(es) da SEAT já tem ausência programada neste período:")
                            for c in choques:
                                st.write(f"- **{c['colaborador']}** — {c['tipo']} de {c['inicio']} a {c['fim']}")

                        st.rerun()
                    else:
                        st.error("Erro ao registrar no banco de dados.")

    # --- ABA 2: QUADRO PÚBLICO ---
    with tab_quadro:
        st.markdown("#### Ausências Programadas, Atestados e Abonos (SEAT)")
        st.caption("Consulte este quadro antes de solicitar férias ou abono para evitar sobreposição de datas na equipe.")

        try:
            todas_ausencias = db_manager.buscar_todos(
                "solicitacoes_ausencia",
                filtros={"setor": "SEAT"},
                ordem_coluna="data_inicio",
                ordem_desc=False,
            ) or []
        except Exception:
            todas_ausencias = []

        # Exibe apenas aprovados e atestados no quadro geral do setor
        publicas = [a for a in todas_ausencias if a.get("status") in ("APROVADA", "NOTIFICADO")]

        if not publicas:
            st.info("Nenhuma ausência ou afastamento programado no momento.")
        else:
            dados_quadro = []
            for a in publicas:
                tipo_raw = a.get("tipo", "AUSENCIA")
                if tipo_raw == "FERIAS":
                    tipo_lbl = "🌴 Férias"
                elif tipo_raw == "ATESTADO":
                    tipo_lbl = "🏥 Atestado"
                elif tipo_raw == "ABONO":
                    tipo_lbl = "📋 Abono"
                else:
                    tipo_lbl = "📝 Ausência"

                ini_str = _formatar_data_curta(a.get("data_inicio"))
                fim_str = _formatar_data_curta(a.get("data_fim"))
                dados_quadro.append({
                    "Colaborador": a.get("colaborador_nome", ""),
                    "Tipo": tipo_lbl,
                    "Período": f"{ini_str} a {fim_str}",
                    "Dias": f"{a.get('dias_afastado', '-')} dia(s)",
                    "Observação": a.get("observacoes", "") or "—",
                })
            df_quadro = pd.DataFrame(dados_quadro)
            st.dataframe(df_quadro, hide_index=True, use_container_width=True)

# ============================================================
# FUNCAO PRINCIPAL
# ============================================================

def renderizar(usuario: dict, modo_edicao: bool = False):
    """Funcao principal do modulo SEAT."""
    import db_manager

    nome = usuario.get("nome", "Usuario")
    cargo = usuario.get("cargo", "operacional")
    setor = usuario.get("setor", "SEAT")

    st.markdown(f"**Colaborador:** {nome} | **Cargo:** {cargo} | **Setor:** {setor}")

    if not modo_edicao:
        st.info("Voce esta em modo de visualizacao. Operacoes de edicao estao bloqueadas.")

    st.markdown("---")

    # Mover Despachos Singulares para Urgentes automaticamente
    _mover_despacho_singular_para_urgentes()

    # Verificar se a sessao foi finalizada
    try:
        todos_processos = db_manager.buscar_todos("pauta_seat") or []
        sessao_finalizada = any(p.get("sessao_finalizada") for p in todos_processos)
    except Exception:
        sessao_finalizada = False

    # Sidebars
    # Despachos Singulares: SEMPRE visivel
    _renderizar_sidebar_ds(usuario)

    # Edicao, Revisao e Urgentes: ocultos quando sessao finalizada
    if not sessao_finalizada:
        _renderizar_sidebar_urgentes(usuario)

    # Escala DOE: SEMPRE visivel
    _renderizar_sidebar_doe(usuario)

    tab_pauta, tab_distribuicao, tab_ds, tab_urgentes, tab_motor, tab_doe, tab_ausencias, tab_gerenciar = st.tabs([
        "Pauta Ativa",
        "Distribuicao",
        "Despachos Singulares",
        "Urgentes",
        "Motor NIP",
        "Escala DOE",
        "Férias e Afastamentos",
        "🗑️ Gerenciar Dados",
    ])
    with tab_pauta:
        _renderizar_pauta_ativa(modo_edicao, usuario)
    with tab_distribuicao:
        _renderizar_distribuicao(modo_edicao, usuario)
    with tab_ds:
        _renderizar_despachos_singulares(modo_edicao, usuario)
    with tab_urgentes:
        _renderizar_urgentes(modo_edicao, usuario)
    with tab_motor:
        _renderizar_motor_nip(modo_edicao, usuario)
    with tab_doe:
        _renderizar_escala_doe(modo_edicao, usuario)
    with tab_ausencias:
        _renderizar_ausencias_seat(modo_edicao, usuario)
    with tab_gerenciar:
        _renderizar_gerenciar_dados(usuario, "SEAT")

def _normalizar_texto(texto):
    """Normaliza texto para comparação (lowercase, sem acentos)."""
    import unicodedata
    if not texto:
        return ""
    texto = str(texto).lower().strip()
    nfkd = unicodedata.normalize('NFKD', texto)
    return ''.join([c for c in nfkd if not unicodedata.combining(c)])

def _renderizar_despachos_singulares(modo_edicao, usuario):
    """Renderiza a tab de Despachos Singulares."""
    st.markdown("### Despachos Singulares")
    st.caption("Processos com despacho singular da sessão atual.")

    try:
        todos_processos = db_manager.buscar_todos("pauta_seat") or []
    except Exception:
        todos_processos = []

    despachos = [
        p for p in todos_processos
        if not p.get("sessao_finalizada", False)
        and not p.get("removido_pauta", False)
        and (
            "despacho singular" in _normalizar_texto(p.get("tipo_sessao", ""))
            or p.get("despacho_singular", False)
        )
    ]

    if not despachos:
        st.info("Nenhum despacho singular na pauta atual.")
        return

    dados = []
    for p in despachos:
        dados.append({
            "Processo": p.get("processo_numero", "—"),
            "Relator": p.get("relator", "—"),
            "Sessão": p.get("numero_sessao", "—"),
            "Dia": str(p.get("dia_sessao", "—"))[:10],
            "Editor": p.get("editor", "—"),
            "Status": "Editado" if p.get("editado") else "Pendente",
        })

    df = pd.DataFrame(dados)
    st.dataframe(df, hide_index=True, use_container_width=True)
